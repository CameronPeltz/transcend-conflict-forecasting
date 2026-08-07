# -*- coding: utf-8 -*-
"""
Rewrites Transcend_DARPA_SBIR_Volume_2_DRAFT.docx into a new file,
paragraph-for-paragraph and table-cell-for-table-cell, preserving every
required heading, DARPA-quoted criterion definition, fact, and number,
while replacing narrative prose with a panel-ready, accessible voice.
Uses substring matching against the ORIGINAL text (not raw paragraph
indices) so a mismatch fails loudly instead of silently writing the
wrong paragraph.
"""
import docx
import sys

SRC = "Transcend_DARPA_SBIR_Volume_2_DRAFT.docx"
DST = "Transcend_DARPA_SBIR_Volume_2_DRAFT_v2.docx"

d = docx.Document(SRC)


def set_text(paragraph, new_text):
    """Replace a paragraph's full text with new_text as a single run,
    preserving the paragraph's style (so Times New Roman / heading
    formatting inherited from the style is untouched)."""
    for r in list(paragraph.runs):
        r.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)


def replace_by_match(paragraphs, needle, new_text, label):
    matches = [p for p in paragraphs if needle in p.text]
    if len(matches) != 1:
        print(f"WARNING [{label}]: found {len(matches)} matches for needle {needle!r} (expected 1)")
        return False
    set_text(matches[0], new_text)
    return True


REPLACEMENTS = [
    # ---------------- Executive Technical Overview ----------------
    ("Transcend (missiontranscend.ai) builds AI decision-support tools",
     "Transcend (missiontranscend.ai) builds AI decision-support tools for peace, security, and geopolitical risk. This proposal extends that mission directly: a conflict early-warning forecasting system built in response to DARPA topic DPA26BZ04-DV015, submitted as a Direct-to-Phase-II Small Business Innovation Research (SBIR) proposal. The topic's central technical requirement is a forecasting mechanism that reasons over a temporal knowledge graph — a live record of who did what to whom, and when, that updates as events unfold — using in-context learning, meaning the system reasons from current evidence and historical analogs at the moment of each forecast rather than being retrained every time new data arrives. That mechanism is not a proposed design: it is built, running, and tested against real, open, publicly available conflict-event data, with the results reported in Part One below."),

    ("Phase I feasibility work, performed by the proposer, produced a working prototype",
     "Phase I feasibility work — performed entirely by the proposer — produced a working prototype tested against three independent, non-overlapping datasets, backtested across more than 1,150 real model configurations, and built on two graph-based model families developed from scratch. The strongest result comes from the Uppsala Conflict Data Program's Georeferenced Event Dataset (UCDP GED), a hand-curated, fatality-coded conflict record maintained independently by Uppsala University and used here as a pure, unmerged comparison set. An alert threshold was chosen using only earlier data, then locked and applied unchanged to a strictly later holdout period the system had never seen. On that honest test, the system reached 84.0% precision at 56.8% recall — meaning that when it flagged a country-week as high risk, it was right 84% of the time, and it caught 57% of the escalations that actually occurred. That result meets the topic's own benchmark for current state-of-the-art precision (approximately 80%) under a validation method designed specifically to prevent the most common way forecasting claims turn out to be false, and it establishes a credible, evidenced path to the Phase II target of 90%."),

    ("Phase II purpose: extend this validated forecasting core",
     "Phase II's purpose is to extend this validated forecasting core with the capability the topic actually exists to fund: multilingual radio ingestion and automatic speech recognition, turning spoken-word radio broadcasts into the same kind of structured signal the system already forecasts from. Rootwise, a radio-listening data aggregation company, joins as a Phase II subcontractor to supply that signal; Applied Peace Strategies contributes international, on-the-ground coordination to prioritize which regions and languages come online first. The forecasting core, graph architecture, and validation discipline built in Phase I do not need to be reinvented for this — they need a new signal connected to them, which is precisely what Phase II funds."),

    ("Dual-use transition potential is direct:",
     "The same forecasting core is not specific to armed conflict by construction, and its transition potential runs in several directions at once: humanitarian early warning, insurance and reinsurance catastrophe modeling, supply-chain risk monitoring, and field-safety planning for journalists and humanitarian organizations working in the same regions this topic names — none of which require export-controlled components. Section 11 details this commercialization path in full."),

    # ---------------- Feasibility intro / metrics primer ----------------
    ("The evidence below is organized one criterion at a time. For each,",
     "The evidence below is organized one criterion at a time, in the order DARPA specifies. For each, this section states what was built, on what real data, how it was tested, what resulted, and — because a proposal that reports only its successes is not trustworthy evidence — where the result falls short today. All code, data-processing scripts, and result files referenced below are maintained in a public repository (github.com/CameronPeltz/transcend-conflict-forecasting) that DARPA evaluators may be granted access to on request. Every dataset used is free and public; no paywalled or classified source was required to reach any result in this section. All validation was performed using rolling-origin backtesting — splitting data by real calendar time, so a model is only ever tested on weeks strictly after the weeks it was trained on — the same discipline financial forecasting and epidemiological forecasting rely on to close the single most common way a forecasting claim turns out to be false: a model that was allowed, even accidentally, to see the future during training.\n\nA handful of terms recur throughout the results that follow, and defining them once here makes every number easier to read without a statistics background. Precision answers the question a decision-maker actually asks: of the country-weeks the system flags as high risk, how many really do escalate? It is the number that determines whether a warning keeps getting read. Recall answers the complementary question: of the escalations that actually happened, how many did the system catch? A useful system needs both — precision without recall is safe but useless, and recall without precision is noisy and gets ignored. Average precision (AP) summarizes ranking quality across every possible alert threshold at once, rather than only the single threshold chosen for a headline result. The Brier score checks whether a stated probability is honest — whether the weeks a system called “70% likely” actually escalate about seven times in ten — on a scale where 0 is a perfect forecaster and roughly 0.25 is what pure guessing produces at a typical base rate. Other metrics used below (specificity, F1, ROC-AUC, log-loss, and the Matthews Correlation Coefficient) are defined at first use and again, with worked examples, in results_v2/final-summary-and-case-studies.html."),

    # ---------------- Criterion 1 ----------------
    ("Method and result. The proposer built and ran a real in-context-learning forecaster",
     "Transcend built and ran a real in-context-learning forecaster (scripts/models.py, class ICLForecaster) that performs no gradient training at prediction time — no weights are ever updated by backpropagation once the system is running, the topic's literal requirement. For each country-week under forecast, the system takes three steps: it serializes the current state of the temporal knowledge graph — recent event volume, the mix of conflict types, tone, and the number of distinct actors involved — into structured text; it retrieves the historical country-weeks most similar to that state through a similarity search over the same real feature space, a retrieval step with no trained model of its own; and it prompts a language model with the current state alongside those retrieved analogs' real, known outcomes, returning both a calibrated probability and a written justification for it."),

    ("Independent evidence of correctness, not just design intent: the pipeline was run",
     "This was not only designed to work — it was run and checked. The full pipeline was executed against a free, locally hosted, open-source language model (Meta's Llama 3.1, run through Ollama, requiring no per-call fee or external API key) across all three of Transcend's datasets, producing genuine model-generated probabilities and reasoning text rather than a scripted fallback. Three real, independently checked examples are documented in results_v2/final-summary-and-case-studies.html (tab 04). In one, the system correctly forecast a one-week-ahead escalation in Colombia (p=0.55, and the escalation occurred) while explicitly discounting its own confidence against the historical base rate for similar weeks. In a second, it correctly forecast a two-week-ahead escalation in Myanmar — and the write-up documents, rather than hides, a real error in the model's own stated reasoning (a misreading of which direction the Goldstein conflict-intensity scale runs). That example is included on purpose: a correct probability and a sound-sounding explanation are not the same kind of evidence, and Phase II's validation work (Part Two, Section 3, Task 5) is designed to keep that distinction visible rather than let a plausible narrative stand in for a checked result."),

    ("Limitation stated plainly: local open-source language model inference currently costs",
     "The honest limitation: local open-source inference currently takes several seconds per prediction, which is why this mechanism was validated on a dedicated, smaller run (results_v2/icl_ollama_log.json) rather than the full 1,150-configuration sweep described under Criterion 2. Closing that latency gap — through model quantization and batched inference, both standard, already-demonstrated techniques that reduce per-call cost without any retraining — is Objective 1 of Phase II (Part Two, Section 2), carried out under Task 3 (Part Two, Section 3)."),

    # ---------------- Criterion 2 ----------------
    ("Method. Rather than report the precision that happens to fall out of one backtest,",
     "Rather than report whatever precision happens to fall out of a single backtest, Transcend ran a validation designed specifically to avoid the most common way a precision claim goes wrong: choosing the alert threshold on the same data used to report performance. scripts/precision_threshold_validation.py splits the UCDP-based forecasting track in time into two windows: an earlier threshold-selection window (August 2021 through August 2023 — 1,627 real predictions, 354 real escalations) and a strictly later, untouched holdout window (August 2023 onward — 858 real predictions, 324 real escalations). A probability threshold was chosen using only the earlier window — the lowest threshold that reached 80% precision there — and then locked. That exact, unchanged threshold was applied to the later window, and only the later window's outcome is reported as the finding."),

    ("Result: 84.0% precision, 56.8% recall, 93.4% specificity, 79.6% overall accuracy,",
     "On 858 real, held-out predictions (324 real escalations), the locked threshold produced 84.0% precision, 56.8% recall, 93.4% specificity (of the country-weeks that did not escalate, 93.4% were correctly left unflagged), and 79.6% overall accuracy, while flagging 25.5% of all country-weeks as high risk. Precision did not erode when the model moved from the window it was tuned on to the window it had never seen — it rose slightly, from 80.3% to 84.0% — real evidence that the threshold generalizes rather than having been fit to the selection window. This clears the topic's own benchmark for current state-of-the-art precision (approximately 80%). The complete precision/recall tradeoff curve, computed honestly on the holdout set alone so evaluators can see what precision is achievable at coverage levels other than the one reported here, is in results_v2/precision_threshold_validation.json."),

    ('What "independently validated" means here, stated precisely rather than left ambiguous:',
     "“Independently validated” means something specific here, worth stating precisely rather than leaving it to imply more than it should. The ground-truth labels come from UCDP, an academic source produced independently of Transcend — that is real independence in the data. The validation method itself (a temporal holdout with a threshold frozen in advance) is as rigorous as an internal validation can be, and it is fully reproducible from the public repository. It has not yet been checked by an evaluator outside Transcend, on data Transcend's team has never seen. That audit is proposed as a funded, explicit Phase II milestone — Month 9, Task 5 (Part Two, Section 3, Table 1) — rather than a claim made as though it had already happened."),

    ("Limitation stated plainly: this result is on the pure UCDP track only.",
     "This result holds on the pure UCDP track. The identical validation run against Transcend's larger, 19-country self-scraped GDELT dataset (the Global Database of Events, Language, and Tone — a free, automated, near-real-time event feed coded from world news) did not reach 80% precision at any useful coverage level; the best achievable there was in the 15–20% precision range at comparable coverage. That is a real negative result, reported rather than left out, and it traces to a specific cause: GDELT's automatically coded event labels carry less signal than UCDP's hand-curated, fatality-coded ones. Closing that gap — and reaching the Phase II 90% target on the broader, novel-signal-enriched data Phase II will collect — is the direct subject of Part Two, Section 2."),

    # ---------------- Criterion 3 ----------------
    ("Current state, stated honestly: this criterion is the least mature of the five",
     "This is the least mature of the five criteria at Phase I, and it is worth being direct about exactly how: no original automatic speech recognition pipeline, radio ingestion system, or synthetic-audio strategy has been built yet. What does exist today is text-based multilingual handling — GDELT's own ingestion pipeline (GDELT Translingual) already translates and codes news coverage in more than 65 languages before it ever reaches Transcend's system, so the forecasting core in this proposal has already been validated end-to-end on signal that originated in many languages, even though Transcend's own code has not yet performed that translation step itself."),

    ("What Phase I did establish, directly relevant to feasibility: the forecasting architecture",
     "What Phase I did establish is directly relevant to closing this gap: the forecasting architecture (Criterion 1) and modeling pipeline (Criterion 2) are agnostic to signal source by construction. Every feature the models consume — event counts, the mix of conflict types, tone, actor and dyad structure — is a property of a coded event record, not of the text-extraction method that produced it. Replacing GDELT- or UCDP-derived events with ASR-transcribed radio events, in whole or in part, is an input-format change rather than an architectural one. That is not a claim made from the whiteboard: it is evidenced by the same model classes in scripts/models.py and scripts/hypergraph_model.py already running unmodified across three structurally different real data sources in this proposal."),

    ("Partnership evidence for closing this gap: Rootwise, a radio-listening data aggregation company,",
     "Two partnerships close the remaining gap. Rootwise, a radio-listening data aggregation company, joins as a Phase II subcontractor specifically for this criterion (Part Two, Sections 3 and 9), bringing an existing radio-monitoring capability Transcend does not need to build from zero. Applied Peace Strategies, through Megan Jeans's international evidence-gathering coordination, brings the ground-level access and language-coverage judgment needed to decide which of the topic's named regions and languages to bring online first. Neither partnership substitutes for the ASR and multimodal engineering work itself, which Part Two, Section 3, Task 2 schedules explicitly against the topic's own Month 3 and Month 6 milestones."),

    # ---------------- Criterion 4 ----------------
    ("Method. Every other backtest in this proposal tests generalization across time",
     "Every other backtest in this proposal tests generalization across time, within countries the model has already trained on — later weeks held out, same countries throughout. That is a real and necessary test, but it is not the harder question Criterion 4 actually asks: does the model still produce a useful forecast in a country it has never seen a single training example from? scripts/loco_validation.py answers that directly with real leave-one-country-out cross-validation. For each of 19 countries in turn, the model trains on the other 18 countries' full real history and is tested only on the one held out — with country identity deliberately excluded as a feature, since a held-out country has no trained coefficient to fall back on and including it would test nothing real."),

    ("Result, pure UCDP track: pooled across all 19 held-out countries,",
     "Pooled across all 19 held-out countries on the pure UCDP track — 28,981 real predictions, 4,814 real escalations — the result is an average precision of 0.617, with 40.7% precision and 71.6% recall. That is meaningfully lower than the same-country result under Criterion 2, which is expected: country identity and country-specific calibration are real, useful signal that this test deliberately removes. But it sits far above the roughly 17% base rate a system with no transferable signal would produce, and it is genuine evidence that the underlying event-volume, conflict-share, and dyad-structure signal transfers across a national border the model never crossed during training. Full per-country results are in results_v2/loco_generalization_validation.json."),

    ("Result, large GDELT track: average precision 0.160 pooled across 19 held-out countries,",
     "The same test on the larger GDELT track produces an average precision of 0.160, pooled across the same 19 held-out countries — close to that track's roughly 12% base rate, which is real but weak generalization. This is consistent with the finding under Criterion 2: it traces again to GDELT's proxy label carrying less transferable signal than UCDP's fatality-coded one, not to any flaw in the leave-one-country-out method itself."),

    ("Limitation stated plainly: this test still only covers 19 countries",
     "This test still covers only the 19 countries Transcend already had event data for. It demonstrates that signal transfers across a border within that set; it does not yet demonstrate performance in a country or region with zero prior event coverage of any kind — a test only Phase II's expanded, radio-sourced collection can make possible for the first time, since it is the first source of signal from places no existing event-coding pipeline reaches."),

    # ---------------- Criterion 5 ----------------
    ("Method and result. The topic names three regions of interest:",
     "The topic names three regions of interest — Central and Southeast Asia, East and Northeast Africa, and South America — and Transcend's larger dataset (three years of self-scraped GDELT across 19 countries) was built deliberately to cover all three: Afghanistan, Myanmar, Pakistan, Tajikistan, Kyrgyzstan, and Uzbekistan for Central/Southeast Asia; Sudan, Ethiopia, Somalia, South Sudan, Kenya, and Eritrea for East/Northeast Africa; and Colombia, Venezuela, Ecuador, Peru, and Bolivia for South America. Seventeen of the nineteen countries sit inside the three named regions; two more, Haiti and Nicaragua, were added outside them specifically to increase the number of real escalation examples available for training, and are disclosed here as exactly that rather than presented as in-scope coverage."),

    ("The pure UCDP comparison track (Criterion 2's strongest result) uses the identical",
     "The pure UCDP comparison track — the source of Criterion 2's headline result — uses this identical 19-country list, which means the 84.0% precision at 56.8% recall reported there is itself evidence of performance inside all three named regions at once, not a single-region result extended by assumption. Per-region and per-country breakdowns are in results_v2/loco_generalization_validation.json and results_v2/final-summary-and-case-studies.html."),

    ("Limitation stated plainly: deployment here means real backtested forecasting performance",
     "“Deployment” here means real, backtested forecasting performance against historical data from these regions — not a live, currently running monitoring feed. That operational step, together with the radio and ASR signal source the topic exists to fund, is Phase II scope (Part Two, Section 3, Task 5)."),

    # ---------------- Existing Technical Foundation ----------------
    ("Prior to this Phase I effort, the proposer produced four internal technical rounds",
     "Before this Phase I effort, Transcend completed four internal technical rounds that establish the modeling and validation discipline this proposal builds on: an initial literature review paired with a 24-configuration hypothesis-driven study (literature-review-and-iterations.html); a corrected model battery, realigned to the topic's actual evaluation criteria after an early design — a temporal graph neural network trained in the ordinary way — was found to violate the topic's no-retraining requirement and replaced before being proposed here (recommended-approach-and-results.html, model-battery-results.html); and a first 1,150-configuration automated search that established the rolling-origin backtesting harness and expanded metrics suite used throughout this proposal (iteration-search-log.html). All four are included in the public repository for evaluator review."),

    ("The working prototype (scripts/ in the public repository) performed by the proposer includes:",
     "The working prototype (scripts/, in the public repository) includes: real data ingestion for GDELT 2.0 and UCDP GED, with data-cleaning steps disclosed rather than hidden (for example, a GDELT retrospective-date-tagging quirk that was found and filtered out rather than silently left in); a temporal knowledge graph construction and feature-engineering layer; the in-context-learning forecaster described under Criterion 1; a hypergraph neural network built from scratch in numpy and scipy (scripts/hypergraph_model.py, hypergraphs_research/), because no hypergraph deep-learning framework was pre-installed in the development environment, and independently cross-checked against xgi, a peer-reviewed open-source hypergraph-analysis library; a graph-based semi-supervised label-spreading classifier; a graph-based natural-language-processing feature pipeline built from a real, 662,355-pair theme co-occurrence graph derived from GDELT's Global Knowledge Graph; and the full rolling-origin backtest harness, which computes ten real metrics — accuracy, precision, recall, specificity, F1 score, Brier score, average precision, ROC-AUC (the odds that the system ranks a true escalation above a false alarm when shown both side by side), log-loss, and the Matthews Correlation Coefficient (a single balanced score, running from −1 to +1, that is not fooled by a lopsided class balance the way plain accuracy can be) — for every configuration tested. Plain-English definitions and worked examples for all ten metrics are provided in results_v2/final-summary-and-case-studies.html for any evaluator without a statistics background."),

    ("results_v2/final-summary-and-case-studies.html is the primary technical report:",
     "results_v2/final-summary-and-case-studies.html is the primary technical report: ranked results across all three datasets and every metric, five cross-track findings, and the three forecasting case studies described under Criterion 1. One finding is included precisely because it is unflattering: 39.7% of label-spreading configurations silently produce zero recall despite scoring competitively on paper — a real operational risk, a system that looks good on a leaderboard while alerting on nothing, reported because DARPA evaluators should see it, not because it helps the pitch. results_v2/iteration-search-log-v2.html is the complete, sortable, per-configuration log behind every summary claim in that report."),

    ("All test data is real and publicly sourced: GDELT 2.0 (data.gdeltproject.org,",
     "All test data is real and publicly sourced. GDELT 2.0 (data.gdeltproject.org, free, no authentication) contributed 5.8 million raw events collected this round; UCDP GED v25.1 (Uppsala University, free, no authentication) contributed 385,918 events across 124 countries worldwide. The model-ready, aggregated country-week panels the models actually train on are committed to the public repository; the largest raw event files are excluded from the repository only for size, and are exactly reproducible from the included, documented download scripts. No proprietary, classified, or paywalled data was used, and none is required to reproduce any result in this proposal."),

    ("results_v2/dangerous-ideas-log.html documents nine radical but explicitly rejected extensions",
     "results_v2/dangerous-ideas-log.html documents nine radical extensions Transcend considered and explicitly rejected during Phase I — among them, scraping private messaging groups for signal, and resolving actor identities down to specific individuals — none implemented, each recorded with the concrete harm it would cause and what a legitimate version would actually require. It is included here because judgment about which capabilities not to build is itself relevant evidence of program-appropriate technical maturity for a dual-use national-security topic."),

    ("None to date. The rolling-origin validation discipline, metrics suite,",
     "None to date. The rolling-origin validation discipline, metrics suite, and hypergraph architecture documented in the public repository are, in Transcend's assessment, of publishable quality; preparing a submission — with DARPA's review, given the topic's sensitivity — is proposed as a Phase II deliverable rather than claimed before it has happened."),

    # ---------------- Commercialization Potential Summary ----------------
    ("Transcend's present team is technical and domain-focused rather than marketing-focused. Commercial go-to-market",
     "Transcend's present team is technical and domain-focused rather than marketing-focused. [Insert: specific advisor, contractor, or hire — name a specific plan here rather than leaving it generic] will bring commercial go-to-market expertise into the company ahead of the Phase II option period, when the transition and commercialization strategy in Section 11 moves from plan to execution."),

    ("The forecasting core validated in this proposal — a temporal-knowledge-graph-based, no-retraining-required",
     "The forecasting core validated in this proposal — a temporal-knowledge-graph forecaster requiring no retraining, with disclosed, rolling-origin-validated precision — is not specific to armed-conflict escalation by construction; the same architecture consumes any event stream with timestamped, geolocated, actor-linked records. Beyond the named topic, Government and DoD applications include humanitarian crisis early warning for USAID and State Department field offices, and force-protection risk indicators for regional combatant commands. Private-sector applications include parametric insurance and reinsurance triggers, where an independently validated, precision-first forecast is directly monetizable; supply-chain and extractives risk monitoring; and field-safety planning for journalists and humanitarian organizations working in the same regions this topic names. Applied Peace Strategies' existing international relationships (Section 11.C) already open a channel into that humanitarian and NGO market — a head start most Phase II companies would otherwise have to build from zero."),

    # ---------------- Part Two, Section 1 ----------------
    ("Conflict early-warning systems today face a precision problem, not an ambition problem.",
     "Conflict early-warning systems today have a precision problem, not an ambition problem. The topic sets the current cross-region state of the art at roughly 80% forecast precision, with a program goal of reaching approximately 90% while cutting within-country false-positive rates in half. Every point of precision below that ceiling carries a real operational cost: an analyst or partner-government official who receives ten warnings and watches seven come true keeps reading the eleventh; one who watches three come true stops. The gap this topic exists to close is not whether a model can output a probability of conflict — many systems already do that — it is whether that probability is trustworthy enough for a real decision-maker to act on, at the lead time and false-alarm rate an operational program actually needs."),

    ("A second, related gap is data-sparse-region coverage.",
     "A second, related gap is coverage of data-sparse regions. The topic names Central and Southeast Asia, East and Northeast Africa, and South America specifically because most existing forecasting systems are built and validated on English-language, internet-dense event streams — and the regions least represented in that kind of data are frequently the ones with the greatest real early-warning need. Radio remains the dominant real-time information channel across much of these regions precisely because it does not depend on the internet infrastructure that a text-based event-coding pipeline like GDELT needs just to see an event at all. A forecasting system validated only on internet-derived text carries a structural blind spot in exactly the places this topic cares about most."),

    ("Transcend addresses both gaps directly. Phase I work (Part One) demonstrates",
     "Transcend addresses both gaps directly, not sequentially. Phase I (Part One) demonstrates that the precision gap is closable with disciplined validation and the right ground-truth data: 84.0% held-out precision on real, fatality-coded conflict data, using the exact mechanism the topic requires — in-context learning over a temporal knowledge graph, with no retraining. Phase II closes the coverage gap, connecting that validated core to Rootwise's radio-listening aggregation capability, with Applied Peace Strategies' international coordination deciding which of the topic's named regions and languages come online first — so the same precision discipline already demonstrated on existing event data extends to the radio-derived signal from data-sparse regions this topic exists to fund."),

    # ---------------- Section 2 Objectives ----------------
    ("Five measurable technical objectives, each mapped directly to one of the five",
     "Five measurable technical objectives follow, each mapped to one of the five Phase I feasibility criteria in Part One and to the Statement of Work tasks in Section 3:"),

    ("Objective 1 — Sustain and harden the no-retraining in-context-learning forecasting mechanism",
     "Objective 1 — Sustain and harden the no-retraining, in-context-learning forecasting mechanism at production latency. Target: reduce per-prediction local-model inference time from several seconds today to under one second, through quantization and batched inference, without degrading the precision and recall results in Objective 2. Validated by Month 6."),

    ("Objective 2 — Reach the topic's 90% cross-region precision target while holding recall",
     "Objective 2 — Reach the topic's 90% cross-region precision target while holding recall at an operationally useful level, and cut within-country false-positive rates in half relative to the Phase I baseline. Target: at least 90% precision on a temporally out-of-sample holdout, with recall at or above 50%, validated on an expanded dataset that combines event-coded and radio-derived signal by Month 9, using the identical frozen-threshold, never-look-ahead method demonstrated in Phase I."),

    ("Objective 3 — Stand up real multilingual, multimodal signal ingestion. Target: working ASR",
     "Objective 3 — Stand up real multilingual, multimodal signal ingestion. Target: working automatic speech recognition for the first tranche of target-region languages by Month 6, with a documented word-error-rate-versus-training-hours curve, and a disclosed real-versus-synthetic training-data ablation by Month 9."),

    ("Objective 4 — Extend generalization evidence from the Phase I leave-one-country-out result",
     "Objective 4 — Extend the Phase I leave-one-country-out generalization result (Part One, Criterion 4) to genuinely new, radio-only-covered locations with no prior GDELT or UCDP event history — proving the architecture's transferability on data it could not have been implicitly tuned against. Target: real backtested performance on at least one radio-sourced-only location by Month 12."),

    ("Objective 5 — Demonstrate the complete, integrated pipeline",
     "Objective 5 — Demonstrate the complete, integrated pipeline — radio ingestion through speech recognition through forecasting through decision-relevant output — on a live or recent real-world scenario, with a quantified ablation showing the novel radio signal's marginal contribution. Target: a Month 12 live demonstration and a Month 18 final benchmark suite covering every objective above, plus the topic's Month 15 paralinguistic-feature extension."),

    # ---------------- Section 3 SOW intro ----------------
    ("This Statement of Work covers an 18-month Phase II base period structured around",
     "This Statement of Work covers an 18-month Phase II base period, organized around the six milestone dates in Table 1, followed by an optional six-month period that converts the demonstrated capability into sustained operational and contingency value. Six tasks carry out the work; each follows the same structure — objective, technical approach, inputs and data, lead and support, dependencies, risks and mitigations, measurable success criteria, deliverables, and start/end month — so progress against every task is checkable the same way."),

    # ---------------- Task 1 ----------------
    ("Task 1 — Radio Collection and Ingestion Engine (Months 1-6, lead: Transcend technical team; support: Rootwise).",
     "Task 1 — Radio Collection and Ingestion Engine (Months 1–6; lead: Transcend technical team; support: Rootwise)."),
    ("Objective: stand up streaming and store-and-forward ingestion from Rootwise's existing radio-listening",
     "Objective: stand up streaming and store-and-forward ingestion from Rootwise's existing radio-listening data aggregation capability for the first tranche of target-region languages, normalized into the same event-record schema the Phase I temporal knowledge graph already consumes."),
    ("Technical approach: Rootwise provides the collection layer (online streaming plus offline software-defined-radio",
     "Technical approach: Rootwise provides the collection layer — online streaming plus offline, software-defined-radio station capture, matching the topic's own described architecture — while Transcend builds the normalization and ingestion adapter that maps Rootwise's raw audio and metadata feed into the graph schema (Actor, Location, Event, Source, and Claim nodes) already validated in Phase I."),
    ("Inputs/data: Rootwise radio feeds; Applied Peace Strategies' regional/language prioritization guidance.",
     "Inputs/data: Rootwise radio feeds; Applied Peace Strategies' regional and language prioritization guidance."),
    ("Dependencies: Rootwise subcontract execution (Section 9); language prioritization from Task 6.",
     "Dependencies: execution of the Rootwise subcontract (Section 9); language prioritization from Task 6."),
    ("Risks/mitigations: radio signal quality in contested or low-infrastructure areas is variable; mitigation is prioritizing station-level",
     "Risks/mitigations: radio signal quality in contested or low-infrastructure areas varies. Mitigation: prioritize station-level redundancy where Rootwise already has multi-station coverage, and pass per-station data-quality flags through to the forecasting layer rather than treating all input as equally reliable."),
    ("Measurable success criteria: ingestion pipeline live for at least 3 target-region languages by Month 3,",
     "Measurable success criteria: ingestion pipeline live for at least three target-region languages by Month 3, with a documented per-language hours-collected count."),
    ("Deliverable: working ingestion engine; Month 3 baseline collection report.",
     "Deliverable: working ingestion engine; Month 3 baseline collection report."),

    # ---------------- Task 2 ----------------
    ("Task 2 — ASR, Language Coverage, and Synthetic-Data Strategy (Months 2-9, lead: Transcend technical team).",
     "Task 2 — Speech Recognition, Language Coverage, and Synthetic-Data Strategy (Months 2–9; lead: Transcend technical team)."),
    ("Objective: fine-tune automatic speech recognition for target-region languages, using a documented,",
     "Objective: fine-tune automatic speech recognition for target-region languages, using a documented, cost-disclosed mix of real annotated audio and synthetic augmentation."),
    ("Technical approach: fine-tune an open, pre-trained speech model (e.g., Whisper Large-v3,",
     "Technical approach: fine-tune an open, pre-trained speech model — for example Whisper Large-v3, which already has strong multilingual baseline performance — for each target language, using native-speaker-annotated audio for the initial real-data tranche. Then run a real ablation across the ratio of real-to-synthetic training audio, to find the most cost-effective mix for each language's available data volume."),
    ("Inputs/data: Rootwise radio audio; native-speaker annotation (budgeted per DoD SBIR guidance at approximately",
     "Inputs/data: Rootwise radio audio; native-speaker annotation, budgeted per DoD SBIR guidance at approximately market rate per annotated hour; synthetic speech generation for underrepresented languages, scoped strictly to training-data augmentation and never distributed or used to impersonate a real, identifiable speaker — the same disclosure standard set in Phase I's dangerous-ideas log."),
    ("Dependencies: Task 1 ingestion; annotation vendor/contractor onboarding.",
     "Dependencies: Task 1 ingestion; annotation vendor and contractor onboarding."),
    ("Risks/mitigations: annotation cost and availability varies significantly by language; mitigation is Applied Peace Strategies'",
     "Risks/mitigations: annotation cost and availability vary significantly by language. Mitigation: Applied Peace Strategies' international coordination identifies native-speaker annotator networks ahead of need, and the real-versus-synthetic ablation directly quantifies how much annotation budget each language actually requires, rather than assuming a fixed rate across all of them."),
    ("Measurable success criteria: word-error-rate-versus-training-hours curve for each onboarded language by Month 6;",
     "Measurable success criteria: a word-error-rate-versus-training-hours curve for each onboarded language by Month 6; a documented real-versus-synthetic ablation by Month 9."),
    ("Deliverables: fine-tuned ASR models; ablation report.",
     "Deliverables: fine-tuned ASR models; ablation report."),

    # ---------------- Task 3 ----------------
    ("Task 3 — Forecasting Integration at Expanded Scale (Months 1-12, lead: Transcend technical team).",
     "Task 3 — Forecasting Integration at Expanded Scale (Months 1–12; lead: Transcend technical team)."),
    ("Objective: extend the validated Phase I forecasting core (in-context-learning forecaster, hypergraph neural network,",
     "Objective: extend the validated Phase I forecasting core — the in-context-learning forecaster, the hypergraph neural network, and graph-based label spreading — to consume the Task 1/2 radio-derived event stream alongside the existing GDELT and UCDP tracks, and re-run the Phase I precision and generalization validation methodology (Part One, Criteria 2 and 4) on the expanded data."),
    ("Technical approach: no forecasting architecture changes are required for new-source integration by design",
     "Technical approach: no architectural changes are required for new-source integration, by design — every Phase I model class consumes a common event-record schema, evidenced by the same code already running unmodified across three structurally different Phase I data sources. This task is a scaling and re-validation effort, not a re-design — and it is also where Objective 1's latency work (quantization and batched inference, reducing per-prediction inference time to under one second) is carried out."),
    ("Inputs/data: Task 1/2 outputs; existing Phase I GDELT/UCDP tracks (retained as a continuity baseline).",
     "Inputs/data: Task 1/2 outputs; the existing Phase I GDELT and UCDP tracks, retained as a continuity baseline."),
    ("Dependencies: Task 1, Task 2.",
     "Dependencies: Task 1; Task 2."),
    ("Risks/mitigations: radio-derived events may have different noise characteristics than text-coded events; mitigation is the same disclosed,",
     "Risks/mitigations: radio-derived events may carry different noise characteristics than text-coded events. Mitigation: the same disclosed, held-out validation discipline used throughout Phase I will surface — and report honestly — any precision degradation, rather than assume the architecture transfers without checking."),
    ("Measurable success criteria: 10-day forecast horizon demonstrated by Month 6; 2-week horizon",
     "Measurable success criteria: a 10-day forecast horizon demonstrated by Month 6; a two-week horizon with precision advancing toward the 90% target by Month 9; a with-versus-without-radio-signal ablation showing a quantified accuracy improvement by Month 12 (topic target: a 30% improvement over the Month 6 baseline while holding precision at the 90% target)."),
    ("Deliverables: expanded forecasting pipeline; Month 9 precision/robustness report;",
     "Deliverables: expanded forecasting pipeline; Month 9 precision and robustness report; Month 12 ablation report and live or recent-scenario demonstration."),

    # ---------------- Task 4 ----------------
    ("Task 4 — Paralinguistic Feature Extraction (Months 12-18, lead: Transcend technical team).",
     "Task 4 — Paralinguistic Feature Extraction (Months 12–18; lead: Transcend technical team)."),
    ("Objective: extend the audio pipeline beyond transcription to prosodic and affective signal",
     "Objective: extend the audio pipeline beyond transcription to prosodic and affective signal — pitch, energy, speaking rate, and indicators of emotional arousal and agitation on the air — and test by ablation whether that signal anticipates events the transcribed words alone do not."),
    ("Technical approach: stand up a paralinguistic feature-extraction layer alongside the Task 2 ASR pipeline",
     "Technical approach: stand up a paralinguistic feature-extraction layer alongside the Task 2 speech-recognition pipeline — a distinct, parallel pipeline over the same audio, not an extension of transcription — and combine these features with the existing textual event stream in the Task 3 forecasting core. Run a direct ablation, forecasting performance with these features against without them, to test the topic's own hypothesis rather than assume it holds."),
    ("Dependencies: Task 1 (audio), Task 2 (aligned transcription for feature-timestamp alignment).",
     "Dependencies: Task 1 (audio); Task 2 (aligned transcription, for feature-timestamp alignment)."),
    ("Risks/mitigations: paralinguistic signal is more culturally and individually variable than lexical content;",
     "Risks/mitigations: paralinguistic signal varies more by culture and by individual than lexical content does. Mitigation: calibrate baselines per language and per region rather than against a single global threshold — consistent with the Phase I finding that country and region identity was one of the strongest real predictive signals found across every dataset tested."),
    ("Measurable success criteria: working paralinguistic feature pipeline across target languages by Month 15;",
     "Measurable success criteria: a working paralinguistic feature pipeline across target languages by Month 15; a documented ablation result — reported honestly whether the signal helps or not — by Month 18."),
    ("Deliverables: paralinguistic feature pipeline; Month 18 final base-period report and benchmark suite.",
     "Deliverables: paralinguistic feature pipeline; Month 18 final base-period report and benchmark suite."),

    # ---------------- Task 5 ----------------
    ("Task 5 — Validation, Independent Audit Readiness, and Live Demonstration (Months 3-18, lead: Transcend technical team; ongoing throughout).",
     "Task 5 — Validation, Independent Audit Readiness, and Live Demonstration (Months 3–18, ongoing; lead: Transcend technical team)."),
    ("Objective: maintain and extend the Phase I validation discipline (rolling-origin backtesting,",
     "Objective: maintain and extend the Phase I validation discipline — rolling-origin backtesting, temporally out-of-sample threshold selection, and leave-one-country-out generalization testing — across every Phase II milestone, and prepare the pipeline for an external evaluation Transcend's own team has not seen in advance."),
    ("Technical approach: every milestone deliverable in this Statement of Work is reported using the same never-look-ahead",
     "Technical approach: every milestone deliverable in this Statement of Work is reported using the same never-look-ahead methodology demonstrated in Part One, published to the same public repository evaluators already have access to. Month 9 specifically includes preparing a held-out evaluation protocol that DARPA, or a DARPA-designated evaluator, can run independently — directly closing the “not yet third-party audited” limitation stated in Part One, Criterion 2."),
    ("Measurable success criteria: Month 12 live or recent-scenario end-to-end demonstration;",
     "Measurable success criteria: a Month 12 live or recent-scenario end-to-end demonstration; a Month 9 independent-audit-ready evaluation protocol delivered."),
    ("Deliverables: mid-program report (Month 12); independent evaluation protocol (Month 9).",
     "Deliverables: mid-program report (Month 12); independent evaluation protocol (Month 9)."),

    # ---------------- Task 6 ----------------
    ("Task 6 — International Evidence-Gathering Coordination (Months 1-18, lead: Applied Peace Strategies, Megan Jeans).",
     "Task 6 — International Evidence-Gathering Coordination (Months 1–18; lead: Applied Peace Strategies, Megan Jeans)."),
    ("Objective: prioritize which target-region languages and locations Task 1 onboards first, based on real,",
     "Objective: prioritize which target-region languages and locations Task 1 onboards first, based on real, current ground-level conditions rather than data-availability convenience alone, and coordinate the native-speaker annotation networks Task 2 depends on."),
    ("Technical approach: Applied Peace Strategies' existing international coordination relationships inform a rolling",
     "Technical approach: Applied Peace Strategies' existing international coordination relationships inform a rolling prioritization schedule, reviewed at each program milestone. This task is coordination and prioritization rather than independent technical development, and is scoped and cost-shared accordingly (Cost Volume)."),
    ("Dependencies: none blocking; informs Tasks 1 and 2 continuously.",
     "Dependencies: none blocking; informs Tasks 1 and 2 continuously."),
    ("Measurable success criteria: documented language/location prioritization delivered ahead of each of Task 1's",
     "Measurable success criteria: documented language and location prioritization delivered ahead of each of Task 1's onboarding phases."),
    ("Deliverables: rolling prioritization memos at Months 3, 6, and 9.",
     "Deliverables: rolling prioritization memos at Months 3, 6, and 9."),

    ("Human/Animal Subjects and/or Recombinant DNA: Not applicable. No research involving human subjects,",
     "Human Subjects, Animal Subjects, and Recombinant DNA: Not applicable. This effort proposes no research involving human subjects, animal subjects, or recombinant DNA."),

    # ---------------- Phase II Option SOW ----------------
    ("Phase II Option Statement of Work (6 months, if exercised). Two directions,",
     "If exercised, the six-month Phase II Option builds directly on base-period deliverables rather than new development, in one of two directions the topic itself frames. The first is an operational pilot and transition: a sustained, live forecasting feed run for a single theater alongside an operational user, measuring real-world warning value over the option period and producing a formal transition package — led by Transcend, with Applied Peace Strategies coordinating the operational-user relationship. The second is rapid language onboarding: standing up a new crisis language in weeks rather than a full collection cycle, using Task 2's real-versus-synthetic ablation findings to minimize new-language annotation cost, and proving a surge capability for contingencies — led by Transcend, with Rootwise providing rapid new-station and new-language collection capacity. Measurable completion criteria and the final deliverable are specified for whichever direction DARPA selects at option exercise; both are evaluated using the identical rolling-origin, never-look-ahead methodology used throughout the base period."),

    # ---------------- Related Work ----------------
    ("Two directly related bodies of prior work, both the proposer's own. First, and most directly:",
     "Two directly related bodies of prior work belong to Transcend. The first, and most directly relevant, is the Phase I feasibility work described in full in Part One — four internal technical rounds (an initial literature review, a corrected model battery, a first 1,150-configuration search, and the feasibility validation reported here), performed entirely by Transcend, documented in the public repository, with no external client or completion date, since this was internally funded feasibility work rather than a delivered engagement. The second, broader in scope, is Transcend's existing AI decision-support platform — strategic decision synthesis, an agentic multi-perspective analysis workflow, and an expert-curated knowledge base, per the company's public description at missiontranscend.ai — which includes the United Nations' first AI agent deployment and reaches the federal market through a Carahsoft contract-vehicle partnership, plus participation in the AWS Activate and NVIDIA Inception programs. That platform is cited here as relevant context for team capability and market access (Section 11), not as Phase I feasibility evidence for the conflict-forecasting-specific criteria in Part One, which stand on their own merits."),

    ("The proposer is aware of the broader state of the art this topic sits within: automated event-coding systems",
     "Transcend is well aware of the broader state of the art this topic sits within: automated event-coding systems (GDELT and ICEWS, the Integrated Crisis Early Warning System) and hand-curated conflict datasets (UCDP and ACLED, the Armed Conflict Location & Event Data Project) as the two dominant ground-truth traditions in computational conflict forecasting; ensemble and machine-learning approaches from the ViEWS research program as a relevant academic precedent for ensemble diversity and rolling-window validation; and graph neural network approaches to event forecasting as an active but — per Transcend's own Phase I literature review — not yet widely operationally deployed research direction, which is exactly why the hypergraph architecture in this proposal was built and empirically tested rather than assumed to work from the literature alone."),

    ("Rootwise's radio-listening data aggregation work and Applied Peace Strategies' international evidence-gathering",
     "Rootwise's radio-listening data aggregation work and Applied Peace Strategies' international evidence-gathering coordination work both predate this proposal and are independent of it. [Insert: a short description of each partner's own directly relevant prior engagement, client/contact, and completion date, per the template's required format, to be supplied by Rootwise and Applied Peace Strategies respectively for Table 3.]"),

    # ---------------- Section 5 ----------------
    ("If successful, Phase II is expected to produce: a forecasting system reaching",
     "If successful, Phase II will produce a forecasting system that reaches the topic's 90% cross-region precision target with the false-positive reduction the topic specifies, validated on real event data spanning coded-text and radio-derived sources across all three named regions; a documented, reusable speech-recognition and synthetic-data methodology for standing up new crisis languages faster than a full from-scratch collection cycle; and a real-versus-synthetic, with-versus-without-radio-signal set of ablation results that settle, with evidence rather than assumption, how much the topic's core novel-signal hypothesis actually contributes to forecast quality."),

    ("Phase II's validated, multi-source forecasting core is the direct foundation for Phase III",
     "Phase II's validated, multi-source forecasting core is the direct foundation for Phase III in two directions already evidenced in this proposal: a DoD or Government operational transition (the Option Period's operational-pilot track, Section 11) and private-sector commercialization (parametric insurance, supply-chain risk, and NGO field-safety applications named in Part One's Commercialization Potential Summary and detailed further in Section 11). Both directions depend on the same underlying asset — an independently auditable, precision-first forecasting pipeline — rather than requiring separate technology development, which is why Phase II is structured to build once and transition twice."),

    # ---------------- Section 6 ----------------
    ("Two key personnel are named for the Phase II effort: Megan Jeans (Applied Peace Strategies)",
     "Two key personnel are named for the Phase II effort: Megan Jeans (Applied Peace Strategies) and David Cyprian (Rootwise), detailed in Table 4 below. [Insert: any additional Transcend technical or management personnel who will be directly involved in Phase II, beyond the personnel already named in Part One and this table.]"),

    # ---------------- Section 9 ----------------
    ("Rootwise (radio-listening data aggregation), point of contact David Cyprian, is proposed as a Phase II",
     "Rootwise (radio-listening data aggregation), point of contact David Cyprian, is proposed as a Phase II subcontractor for Task 1 (radio collection and ingestion), supporting Task 2 (audio for ASR fine-tuning) and the Option Period's rapid-onboarding track. Rootwise's existing radio-monitoring infrastructure directly answers this topic's radio-data-collection-engine requirement — Transcend is not proposing to build station-level radio collection from zero. [Insert: Rootwise's proposed workshare as a percentage of direct and indirect Phase II cost, reconciled with the Cost Volume's Cost Breakdown Structure, and confirmation that at least 50% of total research and analytical work remains with Transcend as the proposing small business, per DoD SBIR Direct-to-Phase-II requirements.]"),

    ("Applied Peace Strategies, point of contact Megan Jeans, coordinates international evidence-gathering (Task 6)",
     "Applied Peace Strategies, point of contact Megan Jeans, coordinates international evidence-gathering (Task 6) as a Transcend subcontractor — prioritizing target-region languages and locations, and supporting native-speaker annotator network access. [Insert: Applied Peace Strategies' proposed workshare as a percentage of direct and indirect Phase II cost, reconciled with the Cost Volume.]"),

    # ---------------- Section 11 ----------------
    ("Transcend is not a pre-revenue concept team: the company's existing AI decision-support platform",
     "Transcend is not a pre-revenue concept team. The company's existing AI decision-support platform — strategic decision synthesis across diplomatic, military, intelligence, economic, and humanitarian data, an agentic multi-perspective analysis workflow, and an expert-curated knowledge base — is already in use, per the company's public description at missiontranscend.ai, including the United Nations' first AI agent deployment, and already reaches the federal market through a Carahsoft contract-vehicle partnership plus participation in AWS Activate and NVIDIA Inception. This Phase I effort is a focused, conflict-early-warning-specific extension of that existing platform and team — not a new company's first product."),

    ("Within that broader context, the conflict-forecasting capability specifically evidenced in Part One is newer:",
     "Within that broader context, the conflict-forecasting capability evidenced specifically in Part One is newer. Its Technology Readiness Level is assessed at TRL 3–4: a validated proof-of-concept with real, disclosed performance metrics (Part One), not yet integrated into Transcend's operational platform or a live monitoring environment. The market for independently validated, precision-first conflict and instability forecasting — distinct from Transcend's broader decision-support product — is early-stage: the closest comparators are academic systems such as ViEWS, which are not packaged as a commercial product, and political-risk consultancies, which do not publish backtested precision or recall figures at all. Phase II milestones that move this capability toward integration with Transcend's existing platform: prototype integration with real radio signal (Months 1–9), independent-audit-ready validation (Month 9), a live or recent-scenario demonstration (Month 12), and a final benchmark suite (Month 18) — full detail in Part Two, Section 3."),

    ("The core problem: organizations that need to act on instability forecasts",
     "Organizations that need to act on instability forecasts — DoD components, humanitarian agencies, insurers, and field-safety teams — currently choose between systems that are fast but unvalidated (a probability number with no disclosed backtest) and systems that are rigorous but academic (published research, not an operational product with a maintained data pipeline). The broader societal need is the same one this topic itself is funding: the gap between having conflict-relevant information and having it early enough, and trusted enough, to act on before a crisis is already underway."),

    ("Primary product: a forecasting feed (API and dashboard) delivering country/region-level escalation probability",
     "The primary product is a forecasting feed — API and dashboard — delivering country- and region-level escalation probability with disclosed precision and recall at the current alert threshold, refreshed as new signal (event-coded today, radio-derived after Phase II) arrives. Primary DoD and federal end-users are regional combatant command intelligence and planning staff, and State Department and USAID field offices that need humanitarian pre-positioning signal. Primary private-sector end-users are parametric insurance and reinsurance underwriters, extractives and supply-chain risk teams, and NGO and journalism field-safety operations — a segment Applied Peace Strategies' existing relationships already serve."),

    ("Primary hypothesis: this forecasting capability is a new module on Transcend's existing go-to-market motion,",
     "This forecasting capability is a new module on Transcend's existing go-to-market motion, not a business model invented from zero. Government and federal access already runs through Transcend's Carahsoft contract-vehicle partnership; the Phase III path — a direct government contract or license for the forecasting feed — uses that same channel. Commercial access across insurance, supply-chain, and NGO segments is a subscription or API product alongside Transcend's existing decision-support offering. The resources needed are a small, forecasting-specific go-to-market function (Part One's Commercialization Potential Summary) and continued partnership with Rootwise and Applied Peace Strategies for signal supply and international coordination, respectively — both already funded under this proposal's cost structure. The core differentiator against academic and commercial-risk-consultancy competitors alike is the disclosed, reproducible, rolling-origin-validated precision figure itself (Part One) — a form of transparency competitors in this space do not currently publish — strengthened by Transcend's existing UN and Carahsoft relationships, which turn DARPA's funding of an independent audit (Task 5) into compounding credibility rather than something to build from scratch."),

    ("Target market: government early-warning/intelligence customers (a small number of large,",
     "The target market runs on two tracks: government early-warning and intelligence customers — a small number of large, high-value contracts, reachable today through Transcend's existing Carahsoft relationship — and a broader commercial base across insurance, extractives and supply-chain risk, and humanitarian and NGO field safety, where Transcend's UN engagement and Applied Peace Strategies' relationships already have a foothold. [Insert: quantified addressable-market sizing — real, cited figures for parametric conflict/political-risk insurance and humanitarian-sector risk-analytics spend, to replace this placeholder before submission.] Competing technologies are political-risk consultancies (qualitative, not backtested), academic systems such as ViEWS (rigorous, but not a maintained commercial product), and a small number of commercial geopolitical-risk-analytics vendors — none of which, to Transcend's knowledge from its own Phase I literature review, publish a rolling-origin-validated precision figure comparable to Part One's. Market validation to date is the Phase I technical validation itself (Part One), together with Transcend's existing UN and Carahsoft relationships and Applied Peace Strategies' relationships (Section 11.C); formal customer discovery specific to the forecasting product is proposed as an early Phase II activity, not claimed as already complete."),

    ("[Insert: Transcend's actual funding history — amount and source of any external financing",
     "[Insert: Transcend's actual funding history — the amount and source of any external financing raised to date — and planned future funding sources (internal revenue, loans, angel, venture capital, or SBIR/STTR as the primary near-term source). This requires real, specific, verifiable company financial information not available to the drafter of this document; do not submit with this placeholder still in place.]"),

    ("Technology risk: the Phase I precision result (84.0% held-out) is real but is on the pure UCDP track only,",
     "Technology risk: the Phase I precision result (84.0% held out) is real, but it is on the pure UCDP track only, not yet on radio-derived signal. Mitigation: Task 3's explicit re-validation on expanded data, using the identical, already-proven methodology — with the honest possibility disclosed here that radio-derived signal could underperform relative to text-coded signal, in which case the Phase II deliverable is that finding itself, reported honestly, rather than a guaranteed precision number. Market risk: the early-stage, fragmented nature of this market (Section 11.E) cuts both ways — little direct competition, but also a customer base not yet trained to expect or ask for disclosed, backtested precision figures. Mitigation: the Option Period's operational-pilot track, which produces exactly the real-world evidence needed to make that case to a first anchor customer. Team risk: the current team's marketing and business-development capacity is limited (Part One, Commercialization Potential Summary). Mitigation: the specific commercial hire or advisor plan named there, in place ahead of the Option Period."),

    ("Transcend's technical team performed all Phase I work described in Part One: the forecasting architecture,",
     "Transcend's technical team performed all Phase I work described in Part One: the forecasting architecture, the graph-native modeling (a hypergraph neural network and label spreading, both built from scratch), the 1,150-configuration validation search, and the temporally out-of-sample precision and generalization testing reported in this proposal. [Insert: specific technical and management team member names, titles, and relevant background beyond the Key Personnel already detailed in Part Two, Section 6, plus an honest statement of financial history and health — cash position, revenue, if any — as the template requires; do not submit without this.] This team has not previously taken a comparable product to market; go-to-market expertise will be added per the plan named in Part One's Commercialization Potential Summary and Section 11.G above."),

    ("[Insert: a real, quantified anticipated-results schedule — additional investment raised, contract value,",
     "[Insert: a real, quantified anticipated-results schedule — additional investment raised, contract value, and/or revenue expected at one year into Phase II, at Phase II completion, and after Phase II completion. DARPA requires this schedule, and requires annual actual-results reporting afterward through the Company Commercialization Report. The drafter of this document has no basis for specific figures here and has intentionally left this as a placeholder rather than invent numbers that would later need to be truthfully reported against.]"),

    # ---------------- Section 12 ----------------
    ("The proposer asserts government purpose rights or better, per standard SBIR data rights protections",
     "Transcend asserts government purpose rights, or better, under standard SBIR data-rights protections (52.227-20), over the original Transcend codebase, model architectures — including the hypergraph neural network and graph-based label-spreading implementations — and trained-model artifacts developed under this effort. No restriction is asserted over the third-party public datasets used (GDELT 2.0 and UCDP GED), which remain governed by their original publishers' public-use terms and are not Transcend intellectual property. [Insert: complete Table 5 below with the specific category, name, basis, and rights entries the template requires, if any additional restricted technical data or computer software applies beyond the above.]"),
]

n_ok, n_fail = 0, 0
for needle, new_text in REPLACEMENTS:
    ok = replace_by_match(d.paragraphs, needle, new_text, needle[:40])
    n_ok += 1 if ok else 0
    n_fail += 0 if ok else 1

print(f"\n{n_ok} paragraph replacements applied, {n_fail} failed/ambiguous")

# ---------------- Table cell rewrites ----------------

def set_cell_text(cell, text):
    cell.text = ""  # clears paragraphs down to one empty paragraph
    p = cell.paragraphs[0]
    p.add_run(text)


tables = d.tables
# Table index 0 = Proposer/Topic/Number (leave as-is)
# Table index 1 = Milestones (Table 2 in the doc's own numbering)
milestone_table = tables[1]
milestone_rows_new = {
    "Month 3": "Radio collection and ingestion stand up for the first target-region languages; native-speaker annotation begins; initial fine-tuned ASR models and the first word-error-rate-versus-training-hours curve are produced; the forecasting baseline is established at the current across-region precision level.",
    "Month 6": "Language coverage expands; the real-to-synthetic training-audio ratio sweep begins; a 10-day forecasting horizon is demonstrated; within-country false-positive rates are cut in half; an accuracy-measurement baseline is established; ASR ablations isolating the radio signal's contribution are reported.",
    "Month 9": "The two-week forecast horizon is demonstrated; across-region precision advances toward the 90% target; the real-versus-synthetic benchmark is delivered; robustness to noise, missing data, and language-coverage gaps is quantified.",
    "Month 12": "An end-to-end demonstration runs on a live or recent real-world scenario; the with-versus-without-radio forecasting ablation is run; a 30% accuracy improvement over the Month 6 baseline is shown while holding precision at the 90% target; a mid-program report is delivered.",
    "Month 15": "The audio model extends beyond transcription to paralinguistic signal: prosodic and affective features — pitch, energy, speaking rate, and indicators of emotional arousal, valence, and agitation — are extracted from broadcast audio; this feature pipeline stands up across target languages; a forecasting baseline combining these features with the textual event stream is established.",
    "Month 18": "An ablation tests whether paralinguistic features add early-warning signal beyond the textual events — whether rising fear, anger, or agitation on the air anticipates events the words alone do not; the final base-period report and benchmark suite is delivered, covering speech recognition, the synthetic-data strategy, textual-event forecasting, and the paralinguistic contribution.",
    "Option Period": "The option period's primary objective is converting the demonstrated capability into operational and contingency value, through one of two directions: an operational pilot and transition, running a sustained live forecasting feed for a single theater alongside an operational user and measuring real-world warning value over six months; or rapid language onboarding, standing up a new crisis language in weeks rather than a full collection cycle, using cross-lingual transfer and synthetic augmentation to prove a surge capability for contingencies.",
}
for row in milestone_table.rows[1:]:
    key = row.cells[0].text.strip()
    if key in milestone_rows_new:
        set_cell_text(row.cells[1], milestone_rows_new[key])
    else:
        print("WARNING: unmatched milestone row key:", repr(key))

# Table index 2 = Related Work (Table 3 in the doc's own numbering)
related_table = tables[2]
if len(related_table.rows) >= 3:
    set_cell_text(related_table.rows[1].cells[1], "Transcend (proposer) — direct foundation for this Phase II proposal")
    set_cell_text(related_table.rows[1].cells[2], "N/A — internally funded feasibility work, no external client")
    set_cell_text(related_table.rows[2].cells[1], "Transcend (proposer) — shares underlying agentic/knowledge-base architecture; broader in scope than this proposal's conflict-forecasting-specific work; cited for team and platform maturity context, not as Phase I feasibility evidence itself")
    set_cell_text(related_table.rows[2].cells[2], "United Nations (per Transcend's public description, its first UN AI agent deployment); Carahsoft (federal contract-vehicle partner) — [Insert Government point of contact and phone/email per template requirement]")

# Table index 3 = Key Personnel (Table 4 in the doc's own numbering)
personnel_table = tables[3]
personnel_quals = {
    "Ola Mohajer": "Former U.S. Institute of Peace and United Nations collaborator; holds overall program leadership. [Insert additional relevant experience and publications.]",
    "Sam Hopkins": "AI systems engineering; technical lead for the forecasting architecture and the hypergraph/in-context-learning implementation (Part Two, Section 3, Tasks 3–5). [Insert additional relevant experience.]",
    "Frank Aum": "Defense and conflict analysis; supports regional and domain validation of forecasting outputs. [Insert additional relevant experience.]",
    "Megan Jeans": "Leads international evidence-gathering coordination (Task 6, Part Two Section 3). [Insert additional relevant technical or domain experience.]",
    "David Cyprian": "Leads radio-listening data aggregation and collection (Task 1, Part Two Section 3); Rootwise point of contact. [Insert additional relevant technical background.]",
}
for row in personnel_table.rows[1:]:
    name_cell_text = row.cells[0].text
    for key, qual in personnel_quals.items():
        if key in name_cell_text:
            set_cell_text(row.cells[3], qual)
            break
    else:
        print("WARNING: unmatched personnel row:", repr(name_cell_text[:40]))

d.save(DST)
print(f"\nSaved {DST}")
