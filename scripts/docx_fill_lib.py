"""
Helper library for filling the DARPA SBIR Volume 2 template in place --
preserves all required DARPA formatting (styles, headings, tables)
rather than rebuilding the document, which is the safer, more compliant
way to produce a submittable file from a mandatory template.
"""
import copy
from docx.shared import Pt
from docx.oxml.ns import qn

BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(11)  # >=10pt per DARPA compliance rule; template's own body default


def find_paragraph_after_heading(doc, heading_text, occurrence=1, placeholder_markers=("[Insert response", "[Response Placeholder", "[Insert required")):
    """Returns the first placeholder-looking OR drafting-guidance-or-normal
    body paragraph found after the Nth (1-indexed) heading paragraph whose
    text contains heading_text. Some sections (e.g. the Statement of Work)
    have a dozen+ Drafting Guidance paragraphs between the heading and the
    real placeholder, so the search window is generous; the caller is
    responsible for forcing correct style/font on whatever is found
    (set_paragraph_text / insert_paragraphs_after both do this already)."""
    seen = 0
    found_heading_idx = None
    paras = doc.paragraphs
    for i, p in enumerate(paras):
        if heading_text in p.text and (p.style.name.startswith("Heading") or p.style.name == "Title" or p.style.name == "Subtitle"):
            seen += 1
            if seen == occurrence:
                found_heading_idx = i
                break
    if found_heading_idx is None:
        raise ValueError(f"heading not found: {heading_text!r} occurrence {occurrence}")
    for j in range(found_heading_idx + 1, min(found_heading_idx + 40, len(paras))):
        p = paras[j]
        if p.style.name.startswith("Heading") or p.style.name in ("Title", "Subtitle"):
            break  # ran into the next section without finding a placeholder
        text = p.text
        if any(m in text for m in placeholder_markers) or text.strip() == "":
            return paras[j], j
    return paras[found_heading_idx + 1], found_heading_idx + 1


def _force_body_format(paragraph, run, bold=False, italic=False):
    paragraph.style = paragraph.part.document.styles["Normal"]
    run.font.name = BODY_FONT
    run.font.size = BODY_SIZE
    run.bold = bold
    run.italic = italic
    # east-asian font field must also be set or Word can silently fall back
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), BODY_FONT)
    rFonts.set(qn('w:hAnsi'), BODY_FONT)
    rFonts.set(qn('w:eastAsia'), BODY_FONT)
    rFonts.set(qn('w:cs'), BODY_FONT)


def set_paragraph_text(paragraph, text, bold=False, italic=False):
    """Replace a paragraph's text entirely with a single run, and force
    it to the document's Normal style / Times New Roman 11pt regardless
    of what style the located placeholder paragraph originally had
    (Drafting Guidance and Response Placeholder are both Arial in this
    template -- non-compliant with the 'Times New Roman' requirement if
    left as-is)."""
    for run in list(paragraph.runs):
        run.text = ""
    if paragraph.runs:
        r = paragraph.runs[0]
    else:
        r = paragraph.add_run()
    r.text = text
    _force_body_format(paragraph, r, bold=bold, italic=italic)
    return paragraph


def insert_paragraphs_after(paragraph, texts, style=None):
    """Insert new paragraphs (each a string) immediately after the given
    paragraph. Always forces Normal style / Times New Roman 11pt on the
    inserted paragraphs regardless of the anchor's original style."""
    from docx.text.paragraph import Paragraph
    anchor = paragraph._p
    new_paragraphs = []
    for t in texts:
        new_p_elm = copy.deepcopy(anchor)
        for r in new_p_elm.findall(qn('w:r')):
            new_p_elm.remove(r)
        anchor.addnext(new_p_elm)
        anchor = new_p_elm
        new_para = Paragraph(new_p_elm, paragraph._parent)
        run = new_para.add_run(t)
        _force_body_format(new_para, run)
        new_paragraphs.append(new_para)
    return new_paragraphs


def fill_section(doc, heading_text, paragraphs_text, occurrence=1, bold_first=False):
    """Find the placeholder paragraph after heading_text and replace it
    with paragraphs_text (a list of strings -> one Word paragraph each)."""
    target, idx = find_paragraph_after_heading(doc, heading_text, occurrence)
    if not paragraphs_text:
        return
    set_paragraph_text(target, paragraphs_text[0], bold=bold_first)
    if len(paragraphs_text) > 1:
        insert_paragraphs_after(target, paragraphs_text[1:])


def remove_drafting_guidance(doc):
    """Strip every 'Drafting Guidance'-styled paragraph -- the template's
    own instruction is to delete these before final submission; they are
    not part of the answer and would otherwise count against the page
    limit for no reason."""
    to_remove = [p for p in doc.paragraphs if p.style.name == "Drafting Guidance"]
    for p in to_remove:
        p._element.getparent().remove(p._element)
    return len(to_remove)


def word_count(doc):
    return sum(len(p.text.split()) for p in doc.paragraphs)


def fill_table_cell(table, row, col, text, bold=False):
    cell = table.cell(row, col)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.name = BODY_FONT
    r.font.size = Pt(10)


def add_table_row(table, texts):
    row = table.add_row()
    for i, t in enumerate(texts):
        if i < len(row.cells):
            row.cells[i].text = str(t)
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.name = BODY_FONT
                    r.font.size = Pt(10)
