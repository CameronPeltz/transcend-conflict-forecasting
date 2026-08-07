# -*- coding: utf-8 -*-
"""
Fills AP_Strategies_DARPA_SBIR_Volume_2_Technical_Template.docx with a
real draft response, in place, preserving the mandatory DARPA template
structure. Every technical claim below traces to a script/result in
this repository (github.com/CameronPeltz/transcend-conflict-forecasting)
or to information the user supplied directly about team/partners.
"""
import sys
sys.path.insert(0, "scripts")
import docx
from docx_fill_lib import fill_section, remove_drafting_guidance, word_count

REPO = "github.com/CameronPeltz/transcend-conflict-forecasting"

# ============================================================ PART ONE

EXEC_SUMMARY = [
    "Transcend (missiontranscend.ai) builds AI decision-support tools for peace, security, and geopolitical risk; this proposal's conflict early-warning forecasting pipeline is a direct extension of that existing mission, built in response to a Direct-to-Phase-II Small Business Innovation Research (SBIR, the federal program this solicitation is issued under; widely used across DoD -- Department of Defense -- proposals) opportunity, DARPA topic DPA26BZ04-DV015. Its core requirement: an in-context-learning (ICL, a way of getting a large language model to produce useful answers by giving it examples and context at the moment of the question, rather than retraining the model itself) forecasting mechanism that runs over a temporal knowledge graph (TKG, a database that tracks not just who-did-what-to-whom but when, so relationships can change over time) without retraining. That mechanism is built, running, and evidenced in this proposal against real, open-source conflict-event data — not a design concept.",
    "Phase I feasibility work, performed by the proposer, produced a working prototype tested across three independent, non-overlapping datasets, more than 1,150 real backtested model configurations, two graph-native model families built from scratch, and a rigorous, temporally out-of-sample validation of forecast precision. On the strongest of the three datasets — the Uppsala Conflict Data Program's Georeferenced Event Dataset (UCDP GED, a real, hand-curated, fatality-coded conflict record maintained by Uppsala University, independent of the proposer), used here as an unmerged, pure comparison set — a threshold chosen only on earlier data and then applied, unchanged, to a strictly later holdout period reached 84.0% precision at 56.8% recall. That meets the topic's own stated ~80% current-state-of-the-art precision benchmark under an honest, disclosed, reproducible validation method, with a credible, evidenced path toward the Phase II 90% target described in Part Two.",
    "Phase II purpose: extend this validated forecasting core with the topic's actual novel-signal requirement — multilingual radio ingestion and automatic speech recognition (ASR, converting spoken audio into text a model can read), delivered in partnership with Rootwise, a radio-listening data aggregation company, and coordinated internationally with support from Applied Peace Strategies for on-the-ground evidence gathering. The forecasting core, graph architecture, and validation discipline built in Phase I do not need to be re-invented for this — they need a new signal source connected to them, which is what Phase II funds.",
    "Dual-use transition potential is direct: the same forecasting core applies to humanitarian early warning, insurance and reinsurance catastrophe modeling, supply-chain risk monitoring, and journalism/non-governmental-organization (NGO, a widely used term for a nonprofit, non-state organization -- in this context, humanitarian and human-rights groups operating in the same regions this topic names) field-safety planning, none of which require export-controlled components. Part Two, Section 11 details the commercialization path.",
]

FEASIBILITY_LEAD_IN = [
    "The evidence below is organized one criterion at a time. For each, this section states the method used, the exact dataset and sample size behind the number, how validation was performed, the result, where the supporting artifact lives, and — per DARPA's own drafting guidance — the limitations plainly. All code, data-processing scripts, and result files cited below are available at " + REPO + ", a repository the proposer controls and can grant DARPA evaluators access to on request. All datasets used are free and public (no paywalled or classified sources were required to reach these results), and all validation was performed by the proposer using rolling-origin backtesting (splitting data by real calendar time, so a model is only ever tested on weeks strictly after the weeks it was trained on) — the same discipline financial forecasting and epidemiological forecasting use to avoid the single most common way forecasting claims turn out to be false: accidentally letting a model see the future during training.",
]

CRITERION_1 = [
    "Method and result. The proposer built and ran a real in-context-learning forecaster (scripts/models.py, class ICLForecaster) that performs no gradient training of the forecasting mechanism at prediction time. For each country-week being forecast, it (1) serializes the current temporal knowledge graph state — recent event volume, conflict-type mix, tone, and distinct-actor counts — to structured text; (2) retrieves the k most similar historical country-weeks by similarity search over that same real feature space (retrieval bookkeeping, not a trained model); and (3) prompts a language model with the current state and the retrieved analogs' real outcomes, returning a calibrated probability and a written justification. The retrieval index is rebuilt from each new training window, but no weights of the forecasting mechanism itself are ever updated by gradient descent — the literal requirement stated in the topic's DP2 evaluation criteria.",
    "Independent evidence of correctness, not just design intent: the pipeline was run against a free, locally hosted, open-source language model (Ollama running Llama 3.1, an open-weight model requiring no per-call fee or external application programming interface (API, the standard way one piece of software requests data or a computation from another) key) across all three of the proposer's datasets, producing genuine model-generated probabilities and reasoning text — not a scripted fallback. Three real, checked-after-the-fact examples are documented in results_v2/final-summary-and-case-studies.html, tab 04, including one where the model correctly forecast an escalation in Colombia one week ahead (p=0.55, escalation occurred) while explicitly discounting its own confidence against the real historical base rate of similar weeks, and one where it correctly forecast a two-week-ahead escalation in Myanmar while the write-up documents, rather than hides, a real reasoning error in the model's own stated justification (a misreading of the Goldstein conflict-intensity scale's direction). That second example is included deliberately: it is honest evidence that a probability output and a plausible-sounding reasoning trace are not interchangeable evidence, a distinction Phase II validation work will formalize.",
    "Limitation stated plainly: local open-source language model inference currently costs several real seconds per prediction, which is why the mechanism was validated on a dedicated smaller run (results_v2/icl_ollama_log.json) rather than across the full 1,150-iteration sweep described under Criterion 2. Phase II Task 4 (Part Two, Section 3) scales this with model quantization and batched inference, both standard, already-demonstrated techniques for reducing per-call latency without retraining.",
]

CRITERION_2 = [
    "Method. Rather than report the precision that happens to fall out of one backtest, the proposer ran a dedicated validation designed specifically to avoid the most common way a precision claim goes wrong: choosing an alert threshold using the same data used to report performance. scripts/precision_threshold_validation.py splits the UCDP-based forecasting track in time into an earlier threshold-selection window (2021-08 through 2023-08, 1,627 real predictions, 354 real escalations) and a strictly later, untouched holdout window (2023-08 onward, 858 real predictions, 324 real escalations). A probability threshold was chosen using only the earlier window, to be the lowest threshold that reached 80% precision there. That exact threshold was then applied, unchanged, to the later window, and only the later window's result is reported as the finding.",
    "Result: 84.0% precision, 56.8% recall, 93.4% specificity, 79.6% overall accuracy, on 858 real, held-out predictions (324 real escalations), flagging 25.5% of all country-weeks as high-risk. Precision did not degrade out of sample — it rose slightly (80.3% to 84.0%) — real evidence the threshold generalizes rather than having been overfit to the selection window. This meets the topic's own stated ~80% current-state-of-the-art precision benchmark. Full results, including the complete precision/recall tradeoff curve computed honestly on the holdout set alone (so DARPA evaluators can see what precision is achievable at other coverage levels, not just the one point reported here), are in results_v2/precision_threshold_validation.json.",
    "What \"independently validated\" means here, stated precisely rather than left ambiguous: the ground-truth labels themselves come from UCDP, an academic source produced independently of the proposer, which is a real form of independence in the data. The validation methodology (temporal holdout, frozen threshold) is designed to be as rigorous as an internal validation can be, and is fully reproducible from the public repository. It has not yet been audited by a third-party evaluator external to the proposer — that audit, using a DARPA-specified or DARPA-collected evaluation set the proposer's team has never seen, is proposed as an explicit, funded Phase II milestone (Month 9, Part Two Section 3/Table 1) rather than a claim made without that audit having happened.",
    "Limitation stated plainly: this result is on the pure UCDP track only. The same validation run against the proposer's larger, 19-country self-scraped GDELT (Global Database of Events, Language, and Tone — a free, automated, near-real-time coded event feed drawn from world news) dataset did not reach 80% precision at any usable coverage level (best achievable was in the 15-20% precision range at comparable coverage) — an honest negative result, reported rather than omitted, and directly attributable to GDELT's automatically-coded event labels carrying less real signal than UCDP's hand-curated, fatality-coded ones. Closing that gap, and reaching DARPA's 90% Phase II target on the broader, novel-signal-enriched data Phase II will collect, is the direct subject of Part Two, Section 2.",
]

CRITERION_3 = [
    "Current state, stated honestly: this criterion is the least mature of the five at Phase I, and the proposer states that plainly rather than overstating it. No original automatic speech recognition (ASR) pipeline, radio ingestion system, or synthetic-audio augmentation strategy has been built yet. What exists today is text-based multilingual signal handling: GDELT's own ingestion pipeline (GDELT Translingual) already translates and codes news coverage in 65-plus languages before it reaches the proposer's system, so the forecasting core in this proposal has already been validated end-to-end on signal that originated in many languages, even though the proposer's own code has not yet performed that translation step itself.",
    "What Phase I did establish, directly relevant to feasibility: the forecasting architecture (temporal knowledge graph plus in-context-learning retrieval, Criterion 1) and the modeling pipeline (Criterion 2) are signal-source-agnostic by construction — every feature the models consume (event counts, conflict-type mix, tone, actor/dyad structure) is a property of a coded event record, not of the text-extraction method that produced it. Swapping in ASR-transcribed radio events in place of (or alongside) GDELT/UCDP-derived events is an input-format change, evidenced by the fact that the identical model classes in scripts/models.py and scripts/hypergraph_model.py already run unmodified across three structurally different real data sources in this proposal.",
    "Partnership evidence for closing this gap: Rootwise, a radio-listening data aggregation company, is named as a Phase II subcontractor specifically for this criterion (Part Two, Sections 3 and 9) — bringing an existing radio-monitoring capability the proposer does not need to build from zero. Applied Peace Strategies' international evidence-gathering coordination (Megan Jeans) provides the ground-level access and language coverage judgment needed to prioritize which of the topic's named regions and languages to onboard first. Neither partnership substitutes for doing the ASR/multimodal engineering work itself, which Part Two, Section 3, Task 2 schedules explicitly against the topic's own Month 3/6 milestones.",
]

CRITERION_4 = [
    "Method. Every other backtest in this proposal tests generalization across time within countries the model has already trained on (rolling-origin: same countries throughout, later weeks held out) — a real and necessary test, but not the harder question DARPA's Criterion 4 actually asks: does the model still produce useful forecasts in a country it has never seen a single training example from? scripts/loco_validation.py answers that directly with real leave-one-country-out (LOCO) cross-validation: for each of 19 countries in turn, the model trains on the other 18 countries' full real history and is tested only on the held-out one, with country identity deliberately excluded as a feature (a held-out country has no trained coefficient to fall back on, so including it would be meaningless).",
    "Result, pure UCDP track: pooled across all 19 held-out countries, 28,981 real predictions, 4,814 real escalations — average precision 0.617, precision 40.7%, recall 71.6%. That is materially lower than the same-country result in Criterion 2 (expected: identity and country-specific calibration are real, useful signal this test deliberately removes), but it is far above the roughly 17% base rate a model with no real transferable signal would produce, and it is genuine evidence that the underlying event-volume, conflict-share, and dyad-structure signal transfers across a national border the model has never crossed in training. Full per-country results are in results_v2/loco_generalization_validation.json.",
    "Result, large GDELT track: average precision 0.160 pooled across 19 held-out countries, close to that track's roughly 12% base rate — real, but weak, generalization. Consistent with Criterion 2's finding, this again traces to GDELT's proxy label carrying less transferable signal than UCDP's fatality-coded one, not to a flaw in the leave-one-country-out method itself.",
    "Limitation stated plainly: this test still only covers 19 countries the proposer already had event data for. It demonstrates that signal transfers across borders within the tested set; it does not yet demonstrate performance in a country or region with zero prior event coverage of any kind, which Phase II's expanded, radio-sourced data collection (novel signal, by the topic's own framing) will make possible to test for the first time.",
]

CRITERION_5 = [
    "Method and result. The topic names three regions of interest: Central and Southeast Asia, East and Northeast Africa, and South America. The proposer's larger dataset (self-scraped GDELT, 3 years, 19 countries) was deliberately built to cover all three: Afghanistan, Myanmar, Pakistan, Tajikistan, Kyrgyzstan, and Uzbekistan for Central/Southeast Asia; Sudan, Ethiopia, Somalia, South Sudan, Kenya, and Eritrea for East/Northeast Africa; and Colombia, Venezuela, Ecuador, Peru, and Bolivia for South America — 17 of the 19 countries sit inside the three named regions. Two additional countries (Haiti and Nicaragua) were added outside the three named regions specifically to increase the number of real escalation examples available for training and are disclosed as such rather than presented as in-scope coverage.",
    "The pure UCDP comparison track (Criterion 2's strongest result) uses the identical 19-country list, so the 84.0% precision / 56.8% recall Phase I headline result is itself evidence of performance inside all three named regions simultaneously, not a single-region result generalized by assertion. Per-region and per-country breakdowns are in results_v2/loco_generalization_validation.json and results_v2/final-summary-and-case-studies.html.",
    "Limitation stated plainly: deployment here means real backtested forecasting performance on historical data from these regions, not a live, currently-running monitoring feed — that operational step, plus the radio/ASR signal source the topic actually funds, is Phase II scope (Part Two, Section 3, Task 5).",
]

FOUNDATION_PREVIOUS_RD = [
    "Prior to this Phase I effort, the proposer produced four internal technical rounds establishing the modeling and validation discipline this proposal builds on: an initial literature review and 24-iteration hypothesis-driven study (literature-review-and-iterations.html); a corrected model battery aligned to the topic's actual DP2 evaluation criteria after identifying that a naive temporal-graph-neural-network design would have violated the topic's no-retraining requirement (recommended-approach-and-results.html, model-battery-results.html); and a first 1,150-iteration automated search establishing the rolling-origin backtesting harness and expanded metrics suite used throughout this proposal (iteration-search-log.html). All are included in the public repository for evaluator review.",
]

FOUNDATION_PROTOTYPE = [
    "The working prototype (scripts/ in the public repository) performed by the proposer includes: real data ingestion for GDELT 2.0 and UCDP GED with disclosed data-cleaning steps (e.g., a documented GDELT retrospective-date-tagging quirk found and filtered rather than silently left in); a temporal knowledge graph construction and feature-engineering layer; the in-context-learning forecaster (Criterion 1); a from-scratch, from-first-principles hypergraph neural network (scripts/hypergraph_model.py, hypergraphs_research/) built in numpy/scipy since no hypergraph deep-learning framework was pre-installed, independently cross-checked using the xgi hypergraph-network-science library (a peer-reviewed open-source package for hypergraph analysis); a graph-based semi-supervised label-spreading classifier; a graph-based natural-language-processing feature pipeline built from a real 662,355-pair theme co-occurrence graph derived from GDELT's Global Knowledge Graph (GKG); and the full rolling-origin backtest harness computing ten real metrics (accuracy, precision, recall, specificity, F1 score, Brier score, average precision, ROC-AUC -- Receiver Operating Characteristic Area Under the Curve, a standard 0-to-1 score for how well a model ranks true positives above false alarms -- log-loss, and Matthews Correlation Coefficient) for every configuration tested. Plain-English definitions and worked examples for every one of these ten metrics are provided in results_v2/final-summary-and-case-studies.html for any evaluator without a statistics background.",
]

FOUNDATION_TECH_REPORTS = [
    "results_v2/final-summary-and-case-studies.html is the primary technical report: ranked results across all three datasets and every metric, five cross-track findings (including a documented, quantified failure mode — 39.7% of label-spreading configurations silently produce zero recall despite competitive ranking scores — reported because a system that scores well on paper while alerting on nothing is a real operational risk, not because it was flattering to include), and the three real forecasting case studies described under Criterion 1. results_v2/iteration-search-log-v2.html is the full, sortable, per-configuration log behind every summary claim.",
]

FOUNDATION_TEST_DATA = [
    "All test data is real and publicly sourced: GDELT 2.0 (data.gdeltproject.org, free, no authentication, 5.8 million raw events collected this round) and UCDP GED v25.1 (Uppsala University, free, no authentication, 385,918 events, 124 countries globally). Derived, model-ready country-week panels (the aggregated tables the models actually train on) are committed to the public repository; the largest raw event files are excluded from the repository for size but are exactly reproducible from the included, documented download scripts. No proprietary, classified, or paywalled data was used or is required to reproduce any result in this proposal.",
]

FOUNDATION_WHITE_PAPERS = [
    "results_v2/dangerous-ideas-log.html documents nine radical but explicitly rejected extensions considered during Phase I feasibility work (e.g., scraping private messaging groups for signal, or resolving actor identities to specific individuals) — none implemented, each with the concrete harm and what a legitimate version would require. Included because a proposer's judgment about which capabilities NOT to build is itself relevant evidence of program-appropriate technical maturity for a dual-use national-security topic.",
]

FOUNDATION_PUBLICATIONS = [
    "None to date. The rolling-origin validation discipline, metrics suite, and hypergraph architecture documented in this proposal's public repository are, in the proposer's assessment, of publishable quality, and preparing a submission (with DARPA's review, given the topic's sensitivity) is proposed as a Phase II deliverable rather than claimed prematurely here.",
]

COMMERCIALIZATION_MARKETING = [
    "Transcend's present team is technical and domain-focused rather than marketing-focused. Commercial go-to-market expertise will be brought in through [Insert: specific advisor, contractor, or hire — the proposer should name a specific plan here rather than leave this generic] ahead of the Phase II option period, when the transition and commercialization strategy in Section 11 moves from planning to execution.",
]

COMMERCIALIZATION_POTENTIAL = [
    "The forecasting core validated in this proposal — a temporal-knowledge-graph-based, no-retraining-required forecasting mechanism with disclosed, rolling-origin-validated precision — is not specific to armed-conflict escalation by construction; the same architecture consumes any event stream with timestamped, geolocated, actor-linked records. Government/DoD applications beyond the named topic include humanitarian crisis early warning for the U.S. Agency for International Development (USAID, the primary federal agency for civilian foreign aid and humanitarian response) and State Department field offices, and force-protection risk indicators for regional combatant commands. Private-sector applications include parametric insurance and reinsurance triggers (where an independently-validated, precision-first forecast is directly monetizable), supply-chain and extractives risk monitoring, and field-safety planning tools for journalism and NGO operations in the same regions this topic already names. Applied Peace Strategies' existing international coordination relationships (Section 11.C) are a direct, already-in-place channel into the humanitarian and NGO segment of this market, ahead of a typical Phase II company having to build that relationship from zero.",
]


# ============================================================ PART TWO

SECTION_1_PROBLEM = [
    "Conflict early-warning systems today face a precision problem, not an ambition problem. The topic states the current cross-region state of the art at roughly 80% forecast precision, with a program goal of reaching approximately 90% while cutting within-country false-positive rates by half. Every point of precision below that ceiling has a direct operational cost: an analyst or a partner government official who receives ten warnings and sees seven materialize will keep reading the eleventh; one who sees three materialize will stop. The gap this topic exists to close is not \"can a model output a probability of conflict\" — many systems already do that — it is \"can a model output a probability of conflict that a real decision-maker can trust enough to act on, at the lead time and false-alarm rate an operational program actually needs.\"",
    "A second, related gap is data-sparse-region coverage. The topic names Central and Southeast Asia, East and Northeast Africa, and South America specifically because existing forecasting systems are disproportionately built and validated on English-language, internet-dense event streams — exactly the regions least represented in that kind of data are frequently the ones with the highest real early-warning need. Radio remains the dominant real-time information channel in many of these regions precisely because it does not require the internet infrastructure that text-based event-coding pipelines like GDELT depend on to see an event at all. A forecasting system validated only on internet-derived text has an structural blind spot in exactly the places this topic cares about most.",
    "Transcend addresses both gaps directly. Phase I work (Part One) demonstrates that the precision gap is closable with disciplined validation and the right ground-truth data — 84.0% held-out precision on real, fatality-coded conflict data, using a forecasting architecture (in-context-learning over a temporal knowledge graph) that matches the topic's own required mechanism exactly. Phase II closes the coverage gap: connecting that validated forecasting core to Rootwise's radio-listening data aggregation capability, with Applied Peace Strategies' international coordination prioritizing which of the topic's named regions and languages to bring online first, so the same precision discipline demonstrated on existing event data extends to the radio-derived, data-sparse-region signal this topic exists to fund.",
]

SECTION_2_OBJECTIVES = [
    "Five measurable technical objectives, each mapped directly to one of the five Phase I feasibility criteria (Part One) and to the Statement of Work tasks in Section 3:",
    "Objective 1 — Sustain and harden the no-retraining in-context-learning forecasting mechanism at production latency. Target: reduce per-prediction local-model inference time from the current several seconds to under one second via quantization and batched inference, without degrading the precision/recall results in Objective 2, validated by Month 6.",
    "Objective 2 — Reach the topic's 90% cross-region precision target while holding recall at an operationally useful level, and cut within-country false-positive rates by half relative to the Phase I baseline. Target: greater than or equal to 90% precision on a temporally out-of-sample holdout, with recall at or above 50%, validated on an expanded, multi-source (event-coded plus radio-derived) dataset by Month 9, using the same frozen-threshold, never-look-ahead validation method demonstrated in Phase I.",
    "Objective 3 — Stand up real multilingual, multimodal signal ingestion. Target: working ASR (automatic speech recognition) for the first tranche of target-region languages by Month 6, with a documented word-error-rate-versus-training-hours curve, and a real, disclosed real-versus-synthetic-training-data ablation by Month 9.",
    "Objective 4 — Extend generalization evidence from the Phase I leave-one-country-out result (Part One, Criterion 4) to genuinely new, radio-only-covered locations with no prior GDELT/UCDP event history, proving the architecture's transferability claim on data it could not have been implicitly tuned against. Target: real backtested performance on at least one radio-sourced-only location by Month 12.",
    "Objective 5 — Demonstrate the complete, integrated pipeline — radio ingestion through ASR through forecasting through decision-relevant output — on a live or recent real-world scenario, with a quantified with-versus-without-radio-signal ablation showing the novel signal source's marginal contribution. Target: Month 12 live demonstration, Month 18 final benchmark suite covering every objective above plus the topic's Month 15 paralinguistic-feature extension.",
]

SECTION_3_INTRO = [
    "This Statement of Work covers an 18-month Phase II base period structured around the six milestone dates below (Table 1), followed by an optional 6-month period converting the demonstrated capability into sustained operational and contingency value. Six tasks organize the work; each is described using the recommended anatomy (objective, technical approach, activities, inputs/data, lead/support, dependencies, risks/mitigations, measurable success criteria, deliverables, start/end month).",
]

TASK_1 = [
    "Task 1 — Radio Collection and Ingestion Engine (Months 1-6, lead: Transcend technical team; support: Rootwise).",
    "Objective: stand up streaming and store-and-forward ingestion from Rootwise's existing radio-listening data aggregation capability for the first tranche of target-region languages, normalized into the same event-record schema the Phase I temporal knowledge graph already consumes.",
    "Technical approach: Rootwise provides the collection layer (online streaming plus offline software-defined-radio station capture, per the topic's own described architecture); Transcend builds the normalization and ingestion adapter mapping Rootwise's raw audio/metadata feed into the existing graph schema (Actor, Location, Event, Source, Claim nodes) validated in Phase I.",
    "Inputs/data: Rootwise radio feeds; Applied Peace Strategies' regional/language prioritization guidance.",
    "Dependencies: Rootwise subcontract execution (Section 9); language prioritization from Task 6.",
    "Risks/mitigations: radio signal quality in contested or low-infrastructure areas is variable; mitigation is prioritizing station-level redundancy where Rootwise has existing multi-station coverage, and disclosing per-station data-quality flags to the forecasting layer rather than silently treating all input as equally reliable.",
    "Measurable success criteria: ingestion pipeline live for at least 3 target-region languages by Month 3, with a documented per-language hours-collected count.",
    "Deliverable: working ingestion engine; Month 3 baseline collection report.",
]

TASK_2 = [
    "Task 2 — ASR, Language Coverage, and Synthetic-Data Strategy (Months 2-9, lead: Transcend technical team).",
    "Objective: fine-tune automatic speech recognition for target-region languages, using a documented, cost-disclosed mix of real annotated audio and synthetic augmentation.",
    "Technical approach: fine-tune an open, pre-trained speech model (e.g., Whisper Large-v3, an existing open-source ASR model with strong multilingual baseline performance) per target language, using native-speaker-annotated audio for the initial real-data tranche, then run a real ablation sweeping the ratio of real-to-synthetic training audio to find the most cost-effective mix for each language's available data volume.",
    "Inputs/data: Rootwise radio audio; native-speaker annotation (budgeted per DoD SBIR guidance at approximately market rate per annotated hour); synthetic speech generation for underrepresented languages, scoped strictly to training-data augmentation and never distributed or used to impersonate a real identifiable speaker, consistent with the disclosure standard set in Phase I's dangerous-ideas log.",
    "Dependencies: Task 1 ingestion; annotation vendor/contractor onboarding.",
    "Risks/mitigations: annotation cost and availability varies significantly by language; mitigation is Applied Peace Strategies' international coordination identifying native-speaker annotator networks ahead of need, and the real-versus-synthetic ablation directly quantifying how much annotation budget is actually necessary per language rather than assuming a fixed rate.",
    "Measurable success criteria: word-error-rate-versus-training-hours curve for each onboarded language by Month 6; documented real-versus-synthetic ablation by Month 9.",
    "Deliverables: fine-tuned ASR models; ablation report.",
]

TASK_3 = [
    "Task 3 — Forecasting Integration at Expanded Scale (Months 1-12, lead: Transcend technical team).",
    "Objective: extend the validated Phase I forecasting core (in-context-learning forecaster, hypergraph neural network, graph-based label spreading) to consume the Task 1/2 radio-derived event stream alongside the existing GDELT/UCDP tracks, and re-run the Phase I precision and generalization validation methodology (Part One, Criteria 2 and 4) on the expanded data.",
    "Technical approach: no forecasting architecture changes are required for new-source integration by design — every Phase I model class consumes a common event-record schema, evidenced by the same code already running unmodified across three structurally different Phase I data sources. This task is primarily a scaling and re-validation effort, not a re-design.",
    "Inputs/data: Task 1/2 outputs; existing Phase I GDELT/UCDP tracks (retained as a continuity baseline).",
    "Dependencies: Task 1, Task 2.",
    "Risks/mitigations: radio-derived events may have different noise characteristics than text-coded events; mitigation is the same disclosed, held-out validation discipline used throughout Phase I, which will surface (and report honestly, per the standard set in Part One) any precision degradation rather than assume the architecture transfers without checking.",
    "Measurable success criteria: 10-day forecast horizon demonstrated by Month 6; 2-week horizon with precision advancing toward the 90% target by Month 9; with-versus-without-radio-signal ablation showing a quantified accuracy improvement by Month 12 (topic target: 30% improvement over the Month 6 baseline while holding precision at the 90% target).",
    "Deliverables: expanded forecasting pipeline; Month 9 precision/robustness report; Month 12 ablation report and live/recent-scenario demonstration.",
]

TASK_4 = [
    "Task 4 — Paralinguistic Feature Extraction (Months 12-18, lead: Transcend technical team).",
    "Objective: extend the audio pipeline beyond transcription to prosodic and affective signal — pitch, energy, speaking rate, and indicators of emotional arousal and agitation on the air — and test by ablation whether that signal anticipates events the transcribed words alone do not.",
    "Technical approach: stand up a paralinguistic feature-extraction layer alongside the Task 2 ASR pipeline (a distinct, parallel pipeline over the same audio, not an extension of transcription); combine these features with the existing textual event stream in the Task 3 forecasting core; run a direct ablation (forecasting performance with versus without paralinguistic features) to test the topic's own stated hypothesis rather than assume it holds.",
    "Dependencies: Task 1 (audio), Task 2 (aligned transcription for feature-timestamp alignment).",
    "Risks/mitigations: paralinguistic signal is more culturally and individually variable than lexical content; mitigation is per-language/per-region baseline calibration rather than a single global threshold, consistent with the Phase I finding that country/region identity was consistently one of the strongest real predictive levers found across every dataset tested.",
    "Measurable success criteria: working paralinguistic feature pipeline across target languages by Month 15; documented ablation result (signal added or not, quantified either way) by Month 18.",
    "Deliverables: paralinguistic feature pipeline; Month 18 final base-period report and benchmark suite.",
]

TASK_5 = [
    "Task 5 — Validation, Independent Audit Readiness, and Live Demonstration (Months 3-18, lead: Transcend technical team; ongoing throughout).",
    "Objective: maintain and extend the Phase I validation discipline (rolling-origin backtesting, temporally out-of-sample threshold selection, leave-one-country-out generalization testing) across every Phase II milestone, and prepare the pipeline for an external, third-party evaluation the proposer's own team has not seen in advance.",
    "Technical approach: every milestone deliverable in this Statement of Work is reported using the same never-look-ahead validation methodology demonstrated in Part One, with results published to the same public repository evaluators already have access to; Month 9 specifically includes preparing a held-out evaluation protocol DARPA (or a DARPA-designated evaluator) can run independently, directly closing the \"not yet third-party audited\" limitation stated plainly in Part One, Criterion 2.",
    "Measurable success criteria: Month 12 live or recent-scenario end-to-end demonstration; Month 9 independent-audit-ready evaluation protocol delivered.",
    "Deliverables: mid-program report (Month 12); independent evaluation protocol (Month 9).",
]

TASK_6 = [
    "Task 6 — International Evidence-Gathering Coordination (Months 1-18, lead: Applied Peace Strategies, Megan Jeans).",
    "Objective: prioritize which target-region languages and locations Task 1 onboards first, based on real, current ground-level conditions rather than data-availability convenience alone, and coordinate the native-speaker annotation networks Task 2 depends on.",
    "Technical approach: Applied Peace Strategies' existing international coordination relationships inform a rolling prioritization schedule reviewed at each program milestone; this task is coordination and prioritization, not independent technical development, and is scoped and cost-shared accordingly (Cost Volume).",
    "Dependencies: none blocking; informs Tasks 1 and 2 continuously.",
    "Measurable success criteria: documented language/location prioritization delivered ahead of each of Task 1's onboarding phases.",
    "Deliverables: rolling prioritization memos at Months 3, 6, and 9.",
]

OPTION_SOW = [
    "Phase II Option Statement of Work (6 months, if exercised). Two directions, per the topic's own framing, both building directly on Phase II base-period deliverables rather than new development: (1) Operational pilot and transition — run a sustained, live forecasting feed for a single theater alongside an operational user for the six-month option period, measuring real-world warning value and producing a formal transition package, led by Transcend with Applied Peace Strategies coordinating the operational-user relationship; (2) Rapid language onboarding — demonstrate standing up a new crisis language in weeks rather than a full collection cycle, using the Task 2 real-versus-synthetic ablation findings to minimize new-language annotation cost, proving a surge capability for contingencies, led by Transcend with Rootwise providing rapid new-station/new-language collection capacity. Measurable completion criteria and final deliverable are specified per whichever direction DARPA selects at option exercise; both are evaluated using the identical rolling-origin, never-look-ahead validation methodology used throughout the base period.",
]

SECTION_4_RELATED_WORK = [
    "Two directly related bodies of prior work, both the proposer's own. First, and most directly: the Phase I feasibility work described fully in Part One -- four internal technical rounds (literature review, corrected model battery, first 1,150-iteration search, and the present feasibility validation) performed entirely by the proposer, documented in the public repository, with no external client or completion date since this was internally funded feasibility work rather than a delivered engagement. Second, and broader in scope: Transcend's existing AI decision-support platform (strategic decision synthesis, agentic multi-perspective analysis, and an expert-curated knowledge base, per the company's public description at missiontranscend.ai), including the United Nations' first AI agent deployment and market access through a Carahsoft federal contract-vehicle partnership plus AWS Activate and NVIDIA Inception program participation -- cited here as directly relevant institutional and technical context for team capability and market access (Section 11), not as Phase I feasibility evidence for the conflict-forecasting-specific criteria in Part One, which is evaluated strictly on its own merits.",
    "The proposer is aware of the broader state of the art this topic sits within: automated event-coding systems (GDELT and ICEWS, the Integrated Crisis Early Warning System -- another automated, machine-coded event dataset with a similar role to GDELT's) and hand-curated conflict datasets (UCDP and ACLED, the Armed Conflict Location & Event Data Project -- a widely used, hand-coded conflict-event dataset comparable to UCDP) as the two dominant ground-truth traditions in computational conflict forecasting; ensemble and machine-learning approaches from the ViEWS (Violence Early-Warning System) research program as a relevant academic precedent for ensemble diversity and rolling-window validation; and graph neural network approaches to event forecasting as an active but, per the proposer's own Phase I literature review, not yet widely operationally deployed research direction — one reason the hypergraph architecture in this proposal was built and empirically tested rather than assumed to work from the literature alone.",
    "Rootwise's radio-listening data aggregation work and Applied Peace Strategies' international evidence-gathering coordination work both predate this proposal and are independent of it; [Insert: a short description of each partner's own directly relevant prior engagement, client/contact, and completion date, per the template's required format, to be supplied by Rootwise and Applied Peace Strategies respectively for Table 2 below].",
]

SECTION_5A = [
    "If successful, Phase II is expected to produce: a forecasting system reaching the topic's 90% cross-region precision target with the false-positive reduction the topic specifies, validated on real event data spanning coded-text and radio-derived sources across all three named regions; a documented, reusable ASR and synthetic-data methodology for standing up new crisis languages faster than a full from-scratch collection cycle; and a real-versus-synthetic and with-versus-without-radio-signal set of ablation results settling, with evidence rather than assumption, how much the topic's core novel-signal hypothesis actually contributes to forecast quality.",
]

SECTION_5B = [
    "Phase II's validated, multi-source forecasting core is the direct foundation for Phase III in two directions already evidenced in this proposal: a DoD/Government operational transition (the Option Period's operational-pilot track, Section 11) and private-sector commercialization (parametric insurance, supply-chain risk, and NGO field-safety applications named in Part One's Commercialization Potential Summary and detailed further in Section 11). Both directions depend on the same underlying asset — an independently-auditable, precision-first forecasting pipeline — rather than requiring separate technology development, which is the basis for treating Phase II as build-once, transition-twice.",
]

FOREIGN_CITIZENS_NOTE = [
    "[REQUIRES PROPOSER CONFIRMATION BEFORE SUBMISSION — do not submit as \"None\" without verifying: this proposal names international coordination (Applied Peace Strategies) and a data-aggregation subcontractor (Rootwise) whose personnel citizenship status is not known to the drafter of this document. If any key personnel, subcontractor staff, or consultant involved in the proposed work are foreign nationals or dual citizens, list each individual's country of origin, visa/work-permit type, and anticipated involvement here, per the template's required format, before submission.]",
]

FACILITIES_EQUIPMENT = [
    "[Insert: proposer's actual facilities description.] Software development and model training for this proposal was performed on standard commercial computing hardware with no specialized laboratory facilities required; no equipment purchase is anticipated beyond standard commercial compute and, if needed, GPU (Graphics Processing Unit, specialized hardware that accelerates the kind of matrix math machine learning models use) instances for ASR fine-tuning, itemized in the Cost Volume. Radio collection hardware and facilities are provided by Rootwise under subcontract (Section 9); [Insert: Rootwise's own facilities/environmental-compliance statement, required by the template for the location where that work is performed].",
]

SUBCONTRACTORS = [
    "Rootwise (radio-listening data aggregation), point of contact David Cyprian, is proposed as a Phase II subcontractor for Task 1 (radio collection and ingestion) and supporting Task 2 (audio for ASR fine-tuning) and the Option Period's rapid-onboarding track. Rootwise's existing radio-monitoring infrastructure is the direct, evidenced answer to this topic's \"radio data collection engine\" requirement — the proposer is not proposing to build station-level radio collection from zero. [Insert: Rootwise's proposed workshare as a percentage of direct and indirect Phase II cost, reconciled with the Cost Volume's Cost Breakdown Structure, and confirmation that at least 50% of total research/analytical work remains with Transcend as the proposing small business, per DoD SBIR Direct-to-Phase-II requirements.]",
    "Applied Peace Strategies, point of contact Megan Jeans, coordinates international evidence-gathering (Task 6) as a Transcend subcontractor: prioritizing target-region languages and locations and supporting native-speaker annotator network access. [Insert: Applied Peace Strategies' proposed workshare as a percentage of direct and indirect Phase II cost, reconciled with the Cost Volume.]",
]

PRIOR_SUPPORT = [
    "No prior, current, or pending support has been provided for proposed work.",
]

DATA_RIGHTS = [
    "The proposer asserts government purpose rights or better, per standard SBIR data rights protections under 52.227-20, over the original Transcend codebase, model architectures (including the hypergraph neural network and graph-based label-spreading implementations), and trained-model artifacts developed under this effort. No restriction is asserted over the third-party public datasets used (GDELT 2.0 and UCDP GED), which remain governed by their original publishers' public-use terms and are not proposer intellectual property. [Insert: complete Table 5 below with the specific category/name/basis/rights entries required by the template if any additional restricted technical data or computer software applies beyond the above.]",
]

# ============================================================ SECTION 11

S11_A = [
    "Transcend is not a pre-revenue concept team: the company's existing AI decision-support platform (strategic decision synthesis across diplomatic, military, intelligence, economic, and humanitarian data; an agentic, multi-perspective analysis workflow; and an expert-curated knowledge base) is already in use, per the company's own public description at missiontranscend.ai, including the United Nations' first AI agent deployment, and reaches the federal market today through a Carahsoft contract-vehicle partnership plus AWS Activate and NVIDIA Inception program participation. This Phase I effort is a focused, conflict-early-warning-specific extension of that existing platform and team, not a new company's first product.",
    "Within that broader context, the conflict-forecasting capability specifically evidenced in Part One is newer: Technology Readiness Level (TRL) for this specific forecasting core is assessed at TRL 3-4 — a validated proof-of-concept with real, disclosed performance metrics (Part One), not yet integrated into Transcend's operational platform or a live monitoring environment. The market for independently-validated, precision-first conflict/instability forecasting specifically (as distinct from Transcend's broader decision-support product) is early-stage: the closest comparators are academic systems (ViEWS) not packaged as a commercial product, and political-risk consultancies that do not publish backtested precision/recall figures at all. Key Phase II milestones advancing this capability toward integration with Transcend's existing platform: prototype integration with real radio signal (Months 1-9), independent-audit-ready validation (Month 9), live/recent-scenario demonstration (Month 12), and a final benchmark suite (Month 18) — see Part Two, Section 3 for full detail.",
]

S11_B = [
    "The core problem: organizations that need to act on instability forecasts — DoD components, humanitarian agencies, insurers, and field-safety teams — currently choose between systems that are fast but unvalidated (a probability number with no disclosed backtest) and systems that are rigorous but academic (published research, not an operational product with a maintained data pipeline). The broader societal need is the same one the topic itself is funding: the gap between having conflict-relevant information and having it early enough, and trusted enough, to act on before a crisis is already underway.",
]

S11_C = [
    "Primary product: a forecasting feed (API and dashboard) delivering country/region-level escalation probability with disclosed precision/recall at the current alert threshold, refreshed as new signal (event-coded and, post-Phase-II, radio-derived) arrives. Primary DoD/Federal end-users: regional combatant command intelligence and planning staff, State Department/USAID field offices needing humanitarian pre-positioning signal. Primary private-sector end-users: parametric insurance and reinsurance underwriters, extractives and supply-chain risk teams, and NGO/journalism field-safety operations already served by Applied Peace Strategies' existing relationships.",
]

S11_D = [
    "Primary hypothesis: this forecasting capability is a new module on Transcend's existing go-to-market motion, not a business model to invent from zero. Government/Federal access already runs through Transcend's Carahsoft contract-vehicle partnership; the Phase III path (direct government contract/licensing for the forecasting feed) uses that existing channel. Commercial access (insurance, supply-chain, NGO segments) is a subscription/API product alongside Transcend's existing decision-support offering. Resources needed: a small commercial go-to-market function specific to the forecasting product (Part One's Commercialization Potential Summary) and continued Rootwise/Applied Peace Strategies partnership for signal supply and international coordination respectively — both already under this proposal's cost structure. Differentiation: the core differentiator against both academic and commercial-risk-consultancy competitors is the disclosed, reproducible, rolling-origin-validated precision figure itself (Part One) — a form of transparency competitors in this space do not currently publish, strengthened by Transcend's existing UN and Carahsoft relationships as credibility and access DARPA's funding of an independent audit (Task 5) would compound rather than have to build from scratch.",
]

S11_E = [
    "Target market: government early-warning/intelligence customers (a small number of large, high-value contracts, reachable today through Transcend's existing Carahsoft federal contract-vehicle relationship) plus a broader commercial base across insurance, extractives/supply-chain risk, and humanitarian/NGO field safety (many smaller subscription customers, a segment Transcend's UN engagement and Applied Peace Strategies' relationships already touch). [Insert: quantified addressable-market sizing — the proposer should supply real, cited market-size figures for parametric conflict/political-risk insurance and for humanitarian-sector risk-analytics spend before submission rather than leave this unquantified]. Competing technologies: political-risk consultancies (qualitative, not backtested); academic systems like ViEWS (rigorous, not a maintained commercial product); a small number of commercial geopolitical-risk-analytics vendors, none of which, to the proposer's knowledge based on Phase I's own literature review, publish a rolling-origin-validated precision figure comparable to Part One's. Market validation to date is the Phase I technical validation itself (Part One) plus Transcend's existing UN/Carahsoft relationships and Applied Peace Strategies' existing relationships (Section 11.C); formal customer discovery specific to the forecasting product is proposed as an early Phase II activity rather than claimed as already complete.",
]

S11_F = [
    "[Insert: Transcend's actual funding history — amount and source of any external financing raised to date — and planned future funding sources (internal revenue, loans, angel, venture capital, or reliance on SBIR/STTR -- Small Business Technology Transfer, a sibling federal non-dilutive funding program to SBIR -- as the primary near-term source). This section requires real, specific, verifiable company financial information the drafter of this document does not have access to; do not submit with this bracketed placeholder still present.]",
]

S11_G = [
    "Technology risk: the Phase I precision result (84.0% held-out) is real but is on the pure UCDP track only, not yet on radio-derived signal; mitigation is Task 3's explicit re-validation on expanded data using the identical, already-proven methodology, with the honest possibility disclosed here that radio-derived signal could underperform relative to text-coded signal, in which case the Phase II deliverable is that finding itself, reported honestly, rather than a guaranteed precision number. Market risk: the early-stage, fragmented nature of this market (Section 11.E) cuts both ways — little direct competition, but also a customer base that has not yet been trained to expect or ask for disclosed, backtested precision figures; mitigation is the Option Period's operational-pilot track, which produces exactly the kind of real-world evidence needed to make that case to a first anchor customer. Team risk: the current team's marketing/business-development capacity is limited (Part One, Commercialization Potential Summary); mitigation is the specific commercial hire/advisor plan to be named there ahead of the Option Period.",
]

S11_H = [
    "Transcend's technical team performed all Phase I work described in Part One: the forecasting architecture, the graph-native modeling (hypergraph neural network and label spreading, both built from scratch), the 1,150-iteration validation search, and the temporally out-of-sample precision and generalization testing in this proposal. [Insert: specific technical/management team member names, titles, and relevant background beyond the Key Personnel already detailed in Part Two, Section 6, plus an honest statement of financial history/health (cash position, revenue, if any) as the template requires — do not submit without this.] This team has not previously taken a comparable product to market; go-to-market expertise will be added per the plan named in Part One's Commercialization Potential Summary and Section 11.G above.",
]

S11_I = [
    "[Insert: a real, quantified anticipated-results schedule — additional investment raised, contract value, and/or revenue expected at one year into Phase II, at Phase II completion, and after Phase II completion. DARPA requires this schedule and requires annual actual-results reporting afterward via the Company Commercialization Report; the drafter of this document has no basis for the specific figures and has intentionally left them as a placeholder rather than inventing numbers that would need to be truthfully reported against later.]",
]


SECTION_6_LEAD_IN = [
    "Two key personnel are named for the Phase II effort: Megan Jeans (Applied Peace Strategies) and David Cyprian (Rootwise), detailed in Table 3 below. [Insert: any additional Transcend technical/management personnel who will be directly involved in Phase II beyond the personnel already named in Part One and this table.]",
]

HUMAN_ANIMAL_SUBJECTS = [
    "Human/Animal Subjects and/or Recombinant DNA: Not applicable. No research involving human subjects, animal subjects, or recombinant DNA is proposed under this effort.",
]


def cleanup_leftovers(doc):
    """Fix specific leftover template artifacts identified during expert
    self-review: orphaned unfilled placeholders the section-heading-based
    fill_section calls didn't reach, a duplicated/contradictory Foreign
    Citizens answer, leftover instructional text duplicated into a
    Response-Placeholder-styled paragraph instead of a real placeholder
    marker, and small pre-existing template typos that would otherwise
    ship into the final submission unchanged."""
    from docx_fill_lib import set_paragraph_text

    paras = doc.paragraphs
    texts = [p.text.strip() for p in paras]

    # A. The two orphaned "[Insert response here.]" placeholders between
    # Task 6 and the Option SOW heading are the template's Human/Animal
    # Subjects sub-slot -- fill with the real, honest "not applicable"
    # statement instead of leaving raw placeholder brackets in a
    # submittable document.
    filled_first = False
    for i, t in enumerate(texts):
        if t in ("[Insert response here.]", "[[Insert response here.]") and not filled_first:
            set_paragraph_text(paras[i], HUMAN_ANIMAL_SUBJECTS[0])
            filled_first = True
        elif t in ("[Insert response here.]", "[[Insert response here.]") and filled_first:
            # second orphan -> remove entirely rather than leave a blank/duplicate
            paras[i]._element.getparent().remove(paras[i]._element)
            filled_first = "done"
            break

    # refresh after structural edit
    paras = doc.paragraphs
    texts = [p.text.strip() for p in paras]

    # B. Table 1 stray caption text
    for p in paras:
        if p.text.strip() == "Table 1. Name X":
            set_paragraph_text(p, "Table 1. Phase II Milestones")
            break

    # D. Foreign Citizens: remove the template's pre-filled "None." --
    # it directly contradicts the caution note this draft inserted right
    # above it, and shipping both is worse than shipping neither.
    for p in paras:
        if p.text.strip() == "None." :
            p._element.getparent().remove(p._element)
            break

    # F. Section 11.H: the template embedded its own drafting-guidance
    # question text inside a "Response Placeholder"-styled paragraph
    # (not a bracketed marker), so the heading-based fill landed on the
    # NEXT paragraph instead and left this instructional text sitting
    # directly above the real answer. Remove it.
    guidance_ghost = ("Expertise/Qualifications of Team/Company Readiness. Describe the "
                       "expertise and qualifications of your management")
    for p in doc.paragraphs:
        if p.text.strip().startswith(guidance_ghost):
            p._element.getparent().remove(p._element)
            break

    # E. Facilities/Equipment: remove the remaining orphaned "[Insert
    # response here.]" placeholder after the real filled paragraph.
    for p in doc.paragraphs:
        if p.text.strip() == "[Insert response here.]":
            p._element.getparent().remove(p._element)
            break

    # C, G: small pre-existing template typos that would otherwise ship
    # unchanged. Text-only fix (does NOT force Normal style / re-run the
    # paragraph) so the "a." item keeps its Heading 2 formatting.
    for p in doc.paragraphs:
        if p.text.strip() == "a. Anticipated technical results.\\" and p.runs:
            p.runs[0].text = p.runs[0].text.replace("results.\\", "results.")
            for extra in p.runs[1:]:
                extra.text = ""
        if "Joint Ethics Regulation, letters from government personnel will NOT be considered" in p.text and p.runs:
            for r in p.runs:
                r.text = r.text.replace(" during the evaluation process.]", " during the evaluation process.")

    # Final comprehensive sweep: remove ANY remaining bare, unfilled
    # "[Insert response here.]"-style orphan the targeted fixes above
    # missed, whatever its root cause -- a submittable document must not
    # contain unresolved template brackets. Genuine bracketed [Insert: ...]
    # notes with real guidance text (which the drafter intentionally left
    # for the proposer to complete with information not available here)
    # are NOT touched by this sweep -- only the bare, contentless markers.
    bare_markers = {"[Insert response here.]", "[[Insert response here.]", "[Insert response hree.]"}
    for p in list(doc.paragraphs):
        if p.text.strip() in bare_markers:
            p._element.getparent().remove(p._element)

    # Font-compliance sweep: the template's own "Response Placeholder" /
    # "Drafting Guidance" styles are Arial, not the required Times New
    # Roman -- any already-answered paragraph in those styles that this
    # script never touched (e.g. Section 10's pre-existing answer) still
    # needs the font forced for template compliance.
    for p in doc.paragraphs:
        if p.style.name == "Response Placeholder" and p.text.strip():
            set_paragraph_text(p, p.text)


def build():
    doc = docx.Document("AP_Strategies_DARPA_SBIR_Volume_2_Technical_Template.docx")

    fill_section(doc, "Executive Technical Overview", EXEC_SUMMARY)
    fill_section(doc, "Demonstration of Phase I Feasibility", FEASIBILITY_LEAD_IN)
    fill_section(doc, "Criterion 1", CRITERION_1)
    fill_section(doc, "Criterion 2", CRITERION_2)
    fill_section(doc, "Criterion 3", CRITERION_3)
    fill_section(doc, "Criterion 4", CRITERION_4)
    fill_section(doc, "Criterion 5", CRITERION_5)
    fill_section(doc, "Previous R&D", FOUNDATION_PREVIOUS_RD)
    fill_section(doc, "Prototype", FOUNDATION_PROTOTYPE)
    fill_section(doc, "Technical reports", FOUNDATION_TECH_REPORTS)
    fill_section(doc, "Test data", FOUNDATION_TEST_DATA)
    fill_section(doc, "White papers", FOUNDATION_WHITE_PAPERS)
    fill_section(doc, "Publications", FOUNDATION_PUBLICATIONS)
    fill_section(doc, "Marketing Expertise", COMMERCIALIZATION_MARKETING)
    fill_section(doc, "Potential for Commercial Application", COMMERCIALIZATION_POTENTIAL)

    # ---- Part Two ----
    fill_section(doc, "1. Identification and Significance", SECTION_1_PROBLEM)
    fill_section(doc, "2. Phase II Technical Objectives", SECTION_2_OBJECTIVES)

    # Section 3: intro + six tasks + option SOW, all under the same "a." placeholder
    sow_all = (SECTION_3_INTRO + [""] + TASK_1 + [""] + TASK_2 + [""] + TASK_3 +
               [""] + TASK_4 + [""] + TASK_5 + [""] + TASK_6)
    fill_section(doc, "3. Phase II Statement of Work", sow_all)
    fill_section(doc, "Phase II Option Statement of Work", OPTION_SOW, occurrence=1)

    fill_section(doc, "4. Related Work", SECTION_4_RELATED_WORK)
    fill_section(doc, "a. Anticipated technical results", SECTION_5A)
    fill_section(doc, "b. Phase II Foundation for Phase III", SECTION_5B)
    fill_section(doc, "6. Key Personnel", SECTION_6_LEAD_IN)
    fill_section(doc, "7. Foreign Citizens", FOREIGN_CITIZENS_NOTE)
    fill_section(doc, "8. Facilities/Equipment", FACILITIES_EQUIPMENT)
    fill_section(doc, "9. Subcontractors/Consultants", SUBCONTRACTORS)
    fill_section(doc, "12. Technical Data Rights", DATA_RIGHTS)

    # ---- Section 11 ----
    fill_section(doc, "A. Transition & Commercialization Summary", S11_A)
    fill_section(doc, "B. Problem / Need Statement", S11_B)
    fill_section(doc, "C. Product / System Applications", S11_C)
    fill_section(doc, "D. Business Model", S11_D)
    fill_section(doc, "E. Target Market", S11_E)
    fill_section(doc, "F. Funding Requirements", S11_F)
    fill_section(doc, "G. Transition & Commercialization Risks", S11_G)
    fill_section(doc, "H. Expertise/Qualifications", S11_H)
    fill_section(doc, "I. Anticipated Transition Results", S11_I)

    # ---- Tables ----
    fill_tables(doc)

    # ---- Cleanup pass (expert self-review findings) ----
    cleanup_leftovers(doc)

    doc.save("Transcend_DARPA_SBIR_Volume_2_DRAFT.docx")
    print("Saved DRAFT. Word count so far:", word_count(doc))


def fill_tables(doc):
    from docx_fill_lib import fill_table_cell, add_table_row

    # Table 0: Proposer cover -- template shipped with an incorrect
    # placeholder proposer name; corrected to the real proposing firm.
    t0 = doc.tables[0]
    fill_table_cell(t0, 1, 0, "Transcend")

    # Table 1: Milestones -- fill Measurable Success Criteria / Deliverable columns (2, 3)
    t1 = doc.tables[1]
    milestone_fill = {
        1: ("Ingestion live for >=3 target-region languages, per-language hours-collected documented.",
            "Working radio ingestion engine; Month 3 baseline collection report."),
        2: ("10-day forecast horizon demonstrated; within-country false-positive rate reduced 50% vs. Phase I baseline; word-error-rate-vs-hours curve documented.",
            "Expanded forecasting pipeline; ASR ablation report."),
        3: ("2-week horizon demonstrated; precision advancing toward 90% target on expanded data; real-vs-synthetic benchmark delivered.",
            "Precision/robustness report; real-vs-synthetic benchmark."),
        4: ("30% accuracy improvement over Month 6 baseline while holding precision at 90% target; with-vs-without-radio ablation quantified.",
            "Live/recent-scenario demonstration; mid-program report."),
        5: ("Paralinguistic feature pipeline live across target languages; combined-feature forecasting baseline established.",
            "Paralinguistic feature pipeline and baseline report."),
        6: ("Paralinguistic-contribution ablation result documented (positive or negative, reported honestly either way).",
            "Final base-period report and benchmark suite (ASR, synthetic-data, forecasting, paralinguistic)."),
    }
    for row_idx, (criteria, deliverable) in milestone_fill.items():
        fill_table_cell(t1, row_idx, 2, criteria)
        fill_table_cell(t1, row_idx, 3, deliverable)
    # Option Period row is the last row
    opt_row = len(t1.rows) - 1
    fill_table_cell(t1, opt_row, 2, "Operational pilot delivering measured real-world warning value OR new crisis language stood up in weeks (per direction selected).")
    fill_table_cell(t1, opt_row, 3, "Transition package OR rapid-onboarding capability report.")

    # Table 2: Related Work
    t2 = doc.tables[2]
    fill_table_cell(t2, 1, 0, "Phase I feasibility validation (this proposal's own Part One work)")
    fill_table_cell(t2, 1, 1, "Transcend (proposer) -- direct foundation for this Phase II proposal")
    fill_table_cell(t2, 1, 2, "N/A -- internally funded feasibility work, no external client")
    fill_table_cell(t2, 1, 3, "2026 (ongoing)")
    add_table_row(t2, [
        "Transcend's existing AI decision-support platform (strategic decision platform, agentic multi-perspective workflow, expert-curated knowledge base) -- publicly described at missiontranscend.ai",
        "Transcend (proposer) -- shared underlying agentic/knowledge-base architecture; broader scope than this proposal's conflict-forecasting-specific work, cited here for team/platform maturity context, not as Phase I feasibility evidence itself",
        "United Nations (per Transcend's own public description, first UN AI agent); Carahsoft (federal contract vehicle partner) -- [Insert Government point of contact and phone/email per template requirement]",
        "Ongoing",
    ])

    # Table 3: Key Personnel -- fill with what's actually known from
    # Transcend's public site (missiontranscend.ai); leave bracketed
    # placeholders for biographical facts (degrees, dates, % commitment)
    # not supplied to the drafter, rather than inventing credentials.
    t3 = doc.tables[3]
    fill_table_cell(t3, 1, 0, "Ola Mohajer -- Founder & CEO [CONFIRM AS PI OR NAME ALTERNATE PI]\n\n[School, Degree, Year -- insert]")
    fill_table_cell(t3, 1, 1, "Transcend")
    fill_table_cell(t3, 1, 2, "[Insert % or hours]")
    fill_table_cell(t3, 1, 3, "Former U.S. Institute of Peace and UN collaborator; overall program leadership. [Insert additional relevant experience/publications.]")
    add_table_row(t3, [
        "Sam Hopkins -- Chief Technology Officer\n\n[School, Degree, Year -- insert]",
        "Transcend",
        "[Insert % or hours]",
        "AI systems engineering; technical lead for forecasting architecture, hypergraph/ICL implementation (Part Two, Section 3, Tasks 3-5). [Insert additional relevant experience.]",
    ])
    add_table_row(t3, [
        "Frank Aum -- Peace & Security Strategist\n\n[School, Degree, Year -- insert]",
        "Transcend",
        "[Insert % or hours]",
        "Defense and conflict analysis; supports regional/domain validation of forecasting outputs. [Insert additional relevant experience.]",
    ])
    add_table_row(t3, [
        "Megan Jeans -- Key Personnel, International Coordination Lead [CONFIRM TITLE]\n\n[School, Degree, Year -- insert]",
        "Applied Peace Strategies (subcontractor)",
        "[Insert % or hours]",
        "International evidence-gathering coordination; leads Task 6 (Part Two, Section 3). [Insert additional relevant technical/domain experience.]",
    ])
    add_table_row(t3, [
        "David Cyprian -- Key Personnel, Radio Data Lead [CONFIRM TITLE]\n\n[School, Degree, Year -- insert]",
        "Rootwise (subcontractor)",
        "[Insert % or hours]",
        "Leads radio-listening data aggregation and collection (Task 1, Part Two Section 3); Rootwise point of contact. [Insert additional relevant technical background.]",
    ])

    # Table 4: Data Rights
    t4 = doc.tables[4]
    add_table_row(t4, [
        "Transcend forecasting codebase, hypergraph and label-spreading model implementations, trained model artifacts",
        "Developed exclusively at private expense under this SBIR effort",
        "SBIR Data Rights (government purpose rights per 52.227-20)",
    ])

    print("Tables filled.")


if __name__ == "__main__":
    build()
